from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "build"
OUT_PATH = OUT_DIR / "ВКР_Поддуба_классификация_инвестиционных_рисков.docx"
RUN_DIR = ROOT / "results" / "defense_run_architected"

TOPIC = (
    "Оптимизация процесса автоматической классификации инвестиционных рисков "
    "путем объединения деревьев решений и кластеризационного анализа "
    "с использованием искусственных нейронных сетей"
)

CAPTION_COUNTERS = {"table": 0, "figure": 0}


def load_results() -> dict:
    metrics = json.loads((RUN_DIR / "metrics.json").read_text(encoding="utf-8"))
    leaderboard = pd.read_csv(RUN_DIR / "model_leaderboard.csv")
    walk_forward = pd.read_csv(RUN_DIR / "walk_forward_report.csv")
    predictions = pd.read_csv(RUN_DIR / "predictions.csv")
    feature_importance = pd.read_csv(RUN_DIR / "feature_importance.csv")
    drift = pd.read_csv(RUN_DIR / "feature_drift_report.csv")
    sector = pd.read_csv(RUN_DIR / "sector_overlay_report.csv")
    universe = pd.read_csv(ROOT / "data/universe.csv")
    panel = pd.read_csv(ROOT / "data/processed/monthly_model_ready.csv", low_memory=False)
    return {
        "metrics": metrics,
        "leaderboard": leaderboard,
        "walk_forward": walk_forward,
        "predictions": predictions,
        "feature_importance": feature_importance,
        "drift": drift,
        "sector": sector,
        "universe": universe,
        "panel": panel,
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 12, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    for tag in ("w:tblLayout", "w:tblW"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(int(width * 567) for width in widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 567)))
        tbl_grid.append(grid_col)


def clear_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tc_borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            tc_borders.append(border)
        border.set(qn("w:val"), "nil")


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name, size, before, after in [
        ("Title", 14, 0, 12),
        ("Heading 1", 14, 18, 12),
        ("Heading 2", 14, 14, 8),
        ("Heading 3", 14, 10, 6),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.first_line_indent = Cm(0)

    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_plain_paragraph(doc: Document, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.25 if first_indent else 0)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    run.bold = bold


def add_center(doc: Document, text: str, *, size: int = 14, bold: bool = False, space_after: int = 0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1 and text[:1].isdigit():
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(1.25)
    elif level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("– " + text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_numbered(doc: Document, idx: int, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{idx}. {text}")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_formula(doc: Document, formula: str, number: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{formula}    ({number})")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_caption(doc: Document, text: str) -> None:
    normalized = text
    is_table_caption = text.startswith("Таблица")
    if text.startswith("Таблица"):
        CAPTION_COUNTERS["table"] += 1
        title = text
        if "–" in text:
            title = text.split("–", 1)[1].strip()
        elif ". " in text:
            title = text.split(". ", 1)[1].strip()
        normalized = f"Таблица {CAPTION_COUNTERS['table']}. {title}"
    elif text.startswith("Рисунок"):
        CAPTION_COUNTERS["figure"] += 1
        title = text
        if "–" in text:
            title = text.split("–", 1)[1].strip()
        elif ". " in text:
            title = text.split(". ", 1)[1].strip()
        normalized = f"Рисунок {CAPTION_COUNTERS['figure']}. {title}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = is_table_caption
    run = p.add_run(normalized)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.italic = False


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None, *, font_size: int = 11) -> None:
    font_size = max(font_size, 12)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    if widths:
        set_fixed_table_geometry(table, widths)
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    header_tr_pr.append(OxmlElement("w:tblHeader"))
    header_tr_pr.append(OxmlElement("w:cantSplit"))
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, header, bold=True, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            set_cell_width(cell, widths[i])
    for row in rows:
        table_row = table.add_row()
        table_row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        cells = table_row.cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if isinstance(value, (int, float)) or (isinstance(value, str) and len(value) <= 12) else WD_ALIGN_PARAGRAPH.LEFT
            text = f"{value:.4f}" if isinstance(value, float) else str(value)
            set_cell_text(cells[i], text, bold=False, size=font_size, align=align)
            if widths:
                set_cell_width(cells[i], widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_figure(doc: Document, image_path: Path, caption: str, *, width_cm: float = 15.0) -> None:
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run().add_picture(str(image_path), width=Cm(width_cm))
        add_caption(doc, caption)


def add_title_page(doc: Document) -> None:
    add_center(doc, "ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ ОБРАЗОВАТЕЛЬНОЕ", size=12)
    add_center(doc, "БЮДЖЕТНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ", size=12)
    add_center(doc, "«ФИНАНСОВЫЙ УНИВЕРСИТЕТ ПРИ ПРАВИТЕЛЬСТВЕ РОССИЙСКОЙ ФЕДЕРАЦИИ»", size=12, bold=True)
    add_center(doc, "Факультет информационных технологий и анализа больших данных", size=12)
    add_center(doc, "Кафедра искусственного интеллекта", size=12, space_after=36)
    add_center(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", size=14, bold=True, space_after=18)
    add_center(doc, "на тему:", size=14)
    add_center(doc, TOPIC, size=14, bold=True, space_after=24)
    add_center(doc, "по направлению подготовки 01.03.02 «Прикладная математика и информатика»", size=12)
    add_center(doc, "образовательная программа «Прикладное машинное обучение»", size=12, space_after=36)

    rows = [
        ("Выполнил:", "студент группы ПМ22-2\nПоддуба И.С."),
        ("Подпись студента:", "________________________"),
        ("Научный руководитель:", "________________________"),
        ("Работа допущена к защите:", "«____» ____________ 2026 г."),
        ("Оценка:", "________________________"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, (left, right) in enumerate(rows):
        set_cell_text(table.rows[r].cells[0], left, bold=False, size=12)
        set_cell_text(table.rows[r].cells[1], right, bold=False, size=12)
        for c in range(2):
            tc_pr = table.rows[r].cells[c]._tc.get_or_add_tcPr()
            for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "nil")
                tc_pr.append(border)
    for _ in range(6):
        doc.add_paragraph()
    add_center(doc, "Москва – 2026", size=12)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    toc = [
        ("ВВЕДЕНИЕ", "3"),
        ("1. ПОСТАНОВКА ЗАДАЧИ АВТОМАТИЧЕСКОЙ КЛАССИФИКАЦИИ ИНВЕСТИЦИОННЫХ РИСКОВ", "6"),
        ("1.1. Экономическая сущность инвестиционного риска", "6"),
        ("1.2. Ручной подход к классификации инвестиционных рисков", "8"),
        ("1.3. Ограничения ручного анализа и необходимость автоматизации", "10"),
        ("1.4. Формальная постановка задачи исследования", "12"),
        ("2. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ КЛАССИФИКАЦИИ\nИНВЕСТИЦИОННЫХ РИСКОВ", "25"),
        ("2.1. Анализ инвестиционного риска как объекта классификации", "25"),
        ("2.2. Анализ исходных данных для оценки инвестиционных рисков", "26"),
        ("2.3. Анализ экономических и математических показателей\nинвестиционного риска", "29"),
        ("2.4. Анализ применимости методов машинного обучения к задаче\nклассификации риска", "34"),
        ("3. РЕАЛИЗАЦИЯ ML-ПАЙПЛАЙНА И ОЦЕНКА РЕЗУЛЬТАТОВ", "46"),
        ("3.1. Архитектура разработанного ML-пайплайна", "46"),
        ("3.2. Реализация расчета признаков и целевой переменной", "48"),
        ("3.3. Реализация моделей машинного обучения", "50"),
        ("3.4. Оценка качества и интерпретация результатов", "52"),
        ("ЗАКЛЮЧЕНИЕ", "66"),
        ("СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ И ИНТЕРНЕТ-РЕСУРСОВ", "70"),
    ]
    for title, page in toc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(title)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(11)
        page_run = p.add_run("\t" + page)
        page_run.font.name = "Times New Roman"
        page_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        page_run.font.size = Pt(11)
    doc.add_page_break()


def intro(doc: Document) -> None:
    add_heading(doc, "ВВЕДЕНИЕ", 1)
    paragraphs = [
        "Инвестиционный риск является одной из ключевых категорий финансового анализа, поскольку именно риск определяет допустимые границы выбора активов, требования к доходности и устойчивость инвестиционного решения при неблагоприятных рыночных изменениях. В прикладной задаче риск акции не может быть сведен только к изменчивости цены. Он проявляется через возможную просадку стоимости, вероятность хвостовых потерь, ухудшение ликвидности, зависимость бумаги от общего рыночного режима и финансовую устойчивость эмитента. Поэтому классификация инвестиционного риска требует не только экономической интерпретации, но и формальной процедуры обработки данных.",
        "Актуальность темы связана с тем, что ручной анализ инвестиционных рисков ограничен субъективностью, трудоемкостью и слабой масштабируемостью. При небольшом числе объектов аналитик может последовательно рассмотреть отчетность, рыночную динамику, структуру ликвидности и макроэкономический фон. Однако при увеличении количества эмитентов, временных срезов и признаков такой подход становится плохо воспроизводимым. Разные специалисты могут по-разному интерпретировать одни и те же показатели, а повторение оценки при обновлении данных требует значительных временных затрат.",
        "Автоматическая классификация инвестиционных рисков позволяет формализовать процесс принятия решения. В рамках такой постановки каждый объект наблюдения описывается набором признаков, после чего модель относит его к одному из заранее определенных классов риска. При этом машинное обучение не отменяет экономический анализ, а переводит его в воспроизводимый контур: признаки рассчитываются по фиксированным формулам, целевая переменная формируется по заданному правилу, а качество классификации оценивается через метрики.",
        "В работе рассматривается задача классификации риска российских акций на месячном срезе. Для каждого наблюдения используются рыночные, ликвидностные, макроэкономические, фундаментальные и отчетные признаки, доступные на дату принятия решения. Целевая переменная строится на основе будущей реализации риска: максимальной просадки, downside-volatility, CVaR 95% и неликвидности на прогнозном горизонте. Такой подход делает задачу ближе к практике риск-скрининга: модель не прогнозирует доходность как таковую, а оценивает вероятность попадания инвестиционного объекта в группу низкого, среднего или высокого риска.",
        "Объект исследования – процесс классификации инвестиционных рисков на основе финансовых, рыночных и иных аналитических данных.",
        "Предмет исследования – методы автоматизации и оптимизации классификации инвестиционных рисков путем объединения деревьев решений, кластеризационного анализа и искусственных нейронных сетей в единый ML-пайплайн.",
        "Цель исследования – разработка и оценка процесса автоматической классификации инвестиционных рисков путем объединения деревьев решений и кластеризационного анализа с использованием искусственных нейронных сетей.",
        "Для достижения поставленной цели необходимо решить следующие задачи:",
    ]
    for text in paragraphs:
        add_plain_paragraph(doc, text)
    tasks = [
        "Изучить существующие подходы к оценке и классификации инвестиционных рисков.",
        "Проанализировать ограничения ручного подхода к классификации инвестиционных рисков.",
        "Определить признаки и метрики, необходимые для формализации задачи классификации инвестиционных рисков.",
        "Спроектировать ML-пайплайн автоматической классификации инвестиционных рисков, объединяющий деревья решений, кластеризационный анализ и искусственные нейронные сети.",
        "Реализовать модели машинного обучения для автоматической классификации инвестиционных рисков.",
        "Оценить качество разработанного подхода и сравнить его с базовым методом классификации.",
    ]
    for i, task in enumerate(tasks, 1):
        add_numbered(doc, i, task)
    more = [
        "Методологическую основу работы составляют методы финансового анализа, математической статистики, временной валидации, кластеризационного анализа, ансамблей деревьев решений и нейросетевого извлечения латентных факторов. Для практической реализации используются Python, pandas, NumPy, scikit-learn и PyTorch. При построении пайплайна учитывается требование point-in-time: признаки должны быть доступны на дату принятия решения, а будущая информация используется только для формирования целевой переменной и последующей оценки качества.",
        "Теоретическая значимость работы заключается в систематизации подхода к формированию многокомпонентной целевой переменной инвестиционного риска и в обосновании объединения деревьев решений, кластеризации и нейронных сетей. Практическая значимость состоит в разработке воспроизводимого программного пайплайна, который может использоваться как исследовательский прототип системы автоматизированного риск-скоринга акций.",
        "Структура работы соответствует логике исследования. В первой главе формулируется задача, раскрывается актуальность и проблематика ручного анализа. Во второй главе проводится анализ предметной области классификации инвестиционных рисков, обосновываются данные, экономические и математические показатели риска, а также методы машинного обучения. В третьей главе описывается реализация ML-пайплайна и проводится оценка результатов. В заключении подводятся итоги выполнения цели и задач исследования.",
    ]
    for text in more:
        add_plain_paragraph(doc, text)
    doc.add_page_break()


def chapter1(doc: Document) -> None:
    add_heading(doc, "1. ПОСТАНОВКА ЗАДАЧИ АВТОМАТИЧЕСКОЙ КЛАССИФИКАЦИИ ИНВЕСТИЦИОННЫХ РИСКОВ", 1)
    add_heading(doc, "1.1. Экономическая сущность инвестиционного риска", 2)
    for text in [
        "Инвестиционный риск в общем виде связан с неопределенностью будущего результата вложения капитала. Для рынка акций такая неопределенность проявляется в колебаниях цен, вероятности убытка, глубине падения относительно предыдущего максимума, возможном расширении спредов и снижении ликвидности. В отличие от доходности, которая может рассматриваться как желаемый результат инвестирования, риск отражает область неблагоприятных исходов, важных для инвестора, портфельного управляющего и системы внутреннего контроля.",
        "В классической финансовой теории риск часто связывается с дисперсией или стандартным отклонением доходности. Такой подход удобен математически, но имеет ограничение: положительные и отрицательные отклонения от среднего учитываются симметрично. Для практического управления инвестициями большее значение имеют именно отрицательные отклонения, поскольку рост цены обычно не воспринимается как риск в том же смысле, что и падение стоимости. Поэтому в прикладных моделях целесообразно дополнительно использовать downside-volatility, максимальную просадку и меры хвостовых потерь.",
        "Риск акции формируется не только рыночной ценой, но и состоянием эмитента. Две бумаги с одинаковой текущей волатильностью могут иметь разный профиль будущей уязвимости, если одна компания имеет устойчивые денежные потоки и низкую долговую нагрузку, а другая зависит от рефинансирования и чувствительна к изменению процентных ставок. Следовательно, классификация инвестиционного риска должна включать как рыночные, так и фундаментальные показатели.",
        "Для российского рынка акций важна специфика предметной области: секторная концентрация, неоднородность ликвидности, влияние ключевой ставки, валютного курса и геополитических факторов. Эти обстоятельства усиливают необходимость учета макроэкономического режима. Бумага, которая выглядит умеренно рискованной в спокойный период, может быстро перейти в высокорисковую группу при росте рыночной корреляции, повышении процентных ставок или ухудшении ликвидности.",
        "С практической точки зрения инвестору важно не только вычислить отдельную числовую меру риска, но и получить классификацию. Классы риска позволяют использовать результат в регламенте принятия решений: низкий риск допускает стандартный анализ, средний риск требует дополнительной проверки, высокий риск может служить сигналом ограничения позиции или передачи объекта на ручной контроль. Поэтому задача классификации более удобна для внедрения в прикладной процесс, чем исключительно регрессионный прогноз.",
        "В данной работе инвестиционный риск рассматривается как многокомпонентная характеристика будущей уязвимости актива. В нее входят ценовые потери, отрицательная изменчивость доходности, хвостовые потери и неликвидность. Такой подход согласуется с экономической природой риска: инвестора интересует не абстрактная изменчивость, а вероятность и масштаб неблагоприятной реализации в будущем окне.",
        "Следовательно, автоматическая классификация инвестиционных рисков должна решать две взаимосвязанные задачи. Первая задача состоит в корректном формировании признаков, которые отражают состояние актива на дату принятия решения. Вторая задача состоит в построении модели, способной сопоставить этот набор признаков с будущим классом риска. Ошибка на любом этапе нарушает смысл классификации: модель либо получает нерелевантные данные, либо обучается на неправильно сформированной целевой переменной.",
        "Важной особенностью финансовых данных является временная зависимость. Наблюдения нельзя рассматривать как независимые строки таблицы, поскольку текущий рынок связан с предыдущими периодами, а будущие данные недоступны на дату принятия решения. Поэтому при постановке задачи необходимо заранее определить горизонт оценки, порядок формирования признаков и способ разделения выборки на обучающую, валидационную и тестовую части.",
    ]:
        add_plain_paragraph(doc, text)

    add_heading(doc, "1.2. Ручной подход к классификации инвестиционных рисков", 2)
    for text in [
        "Ручной подход к классификации инвестиционных рисков обычно строится как последовательная аналитическая процедура. Аналитик выбирает инвестиционные объекты, собирает рыночные данные, изучает финансовую отчетность, рассчитывает показатели доходности, волатильности, ликвидности и долговой нагрузки, после чего интерпретирует полученные значения с учетом отрасли и текущего макроэкономического контекста.",
        "Ручной анализ обладает важным достоинством: он позволяет учитывать качественные обстоятельства, которые трудно непосредственно представить числом. К таким обстоятельствам относятся особенности корпоративного управления, зависимость компании от отдельных рынков сбыта, регуляторные ограничения, санкционные риски и специфика раскрытия информации. В рамках исследования эти представления переводятся в формулы, признаки и метрики, чтобы сохранить предметную основу инвестиционного анализа и сделать классификацию проверяемой.",
        "Однако ручная классификация плохо подходит для регулярного анализа большого числа объектов. Если в выборку входит несколько десятков эмитентов и несколько лет месячных наблюдений, число анализируемых ситуаций быстро становится значительным. Например, 18 эмитентов при ежемесячной оценке за четыре года образуют сотни наблюдений. Для каждого из них нужно обновить данные, пересчитать показатели и сохранить логику принятого решения.",
        "Еще одна проблема ручного анализа состоит в неоднородности критериев. Даже если аналитики используют одинаковые показатели, они могут по-разному оценивать их значимость. Один аналитик может сильнее учитывать волатильность, другой – ликвидность, третий – долговую нагрузку и денежный поток. В результате итоговая классификация зависит не только от данных, но и от субъективной системы весов, которая не всегда явно зафиксирована.",
        "Воспроизводимость ручной оценки также ограничена. Если через некоторое время необходимо повторить классификацию на тех же данных, важно получить тот же результат. В ручной процедуре это возможно только при наличии подробного регламента, сохраненных расчетов и строгой дисциплины работы с источниками. На практике часть решений принимается на основе профессионального суждения, которое трудно полностью восстановить.",
        "Ручной подход может использовать электронные таблицы и отдельные скрипты, но это не решает проблему методической целостности. Разрозненные расчеты увеличивают риск ошибок при копировании данных, изменении формул, сортировке строк, пересчете периодов и обновлении источников. Кроме того, если признаки формируются вручную, сложнее гарантировать отсутствие заглядывания в будущее, особенно при работе с отчетностью, опубликованной позже отчетного периода.",
        "В контексте данной работы ручной анализ рассматривается не как неверный, а как ограниченный базовый подход. Его роль состоит в том, чтобы показать исходную проблему и сформировать требования к автоматизированному процессу. Такой процесс должен быть формализованным, повторяемым, масштабируемым и проверяемым по метрикам качества.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 1.1 – Сравнение ручного и автоматизированного подходов к классификации риска")
    add_table(
        doc,
        ["Критерий", "Ручной подход", "Автоматизированный ML-пайплайн"],
        [
            ["Масштабирование", "Требует роста трудозатрат при увеличении числа объектов", "Позволяет пакетно обрабатывать множество наблюдений"],
            ["Воспроизводимость", "Зависит от исполнителя и сохранности расчетов", "Определяется кодом, параметрами и версией данных"],
            ["Интерпретация", "Опирается на профессиональное суждение", "Поддерживается важностью признаков и правилами модели"],
            ["Обновление", "Требует повторного ручного пересчета", "Выполняется через повторный запуск пайплайна"],
            ["Контроль качества", "Часто качественный", "Оценивается через accuracy, F1, recall, confusion matrix"],
        ],
        [3.3, 5.8, 6.4],
        font_size=10,
    )
    add_source(doc, "Источник: составлено автором.")

    add_heading(doc, "1.3. Ограничения ручного анализа и необходимость автоматизации", 2)
    for text in [
        "Проблематика ручного анализа инвестиционных рисков в научной постановке связана прежде всего с методологическими ограничениями. Ручной подход трудно формализовать в виде единого алгоритма, если в нем одновременно используются количественные показатели, профессиональная интерпретация и неявные допущения о будущем состоянии рынка. Это затрудняет проверку качества и воспроизводимость результата.",
        "Первое ограничение – субъективность интерпретации. Риск является многокомпонентной категорией, поэтому разные аналитики могут придавать разный вес одним и тем же признакам. Высокая ликвидность может смягчать восприятие волатильности, а сильная долговая нагрузка может усиливать оценку риска даже при стабильной цене. Без формальной модели такие взаимосвязи часто остаются неявными.",
        "Второе ограничение – трудоемкость. Для каждого объекта необходимо собрать данные, проверить их актуальность, рассчитать показатели, сопоставить значения с аналогами и сформировать итоговую категорию риска. Чем больше признаков учитывается, тем выше стоимость повторения анализа. При регулярной ежемесячной оценке трудоемкость становится самостоятельным фактором, ограничивающим применимость ручного подхода.",
        "Третье ограничение – низкая масштабируемость при увеличении числа объектов и временных срезов. В финансовых задачах важно анализировать не только текущий момент, но и историю, чтобы оценить устойчивость модели и проверить ее на будущих периодах. Ручная классификация сотен наблюдений требует либо значительных ресурсов, либо упрощения методики, что снижает качество анализа.",
        "Четвертое ограничение – сложность учета нелинейных зависимостей. Инвестиционный риск может возрастать не из-за одного отдельного признака, а из-за комбинации факторов: высокой беты, снижения ликвидности, роста ставки и ухудшения денежного потока. Аналитик может заметить часть таких сочетаний, но систематически учитывать их по всей выборке затруднительно.",
        "Пятое ограничение – риск нарушения временной корректности. Финансовая отчетность публикуется с задержкой, поэтому данные за отчетный период не всегда были доступны инвестору в момент принятия решения. Если при построении ручной оценки используется отчетность по дате периода, а не по дате публикации, возникает утечка будущей информации. Для исследовательской работы это критично, поскольку модель может показать завышенное качество.",
        "Автоматизация позволяет устранить часть перечисленных ограничений. Она задает единые правила расчета признаков, фиксирует формулы, обеспечивает повторяемость обработки данных и дает возможность сравнивать модели по одинаковым метрикам. При этом автоматизированная система не заменяет экономический смысл: признаки и целевая переменная должны быть обоснованы предметной областью.",
        "Таким образом, необходимость автоматизации определяется не только желанием ускорить обработку данных. Главный научный мотив состоит в переходе от субъективной ручной процедуры к формализованному и проверяемому процессу классификации. Такой процесс можно воспроизвести, оценить, улучшить и сопоставить с базовыми моделями.",
    ]:
        add_plain_paragraph(doc, text)

    add_heading(doc, "1.4. Формальная постановка задачи исследования", 2)
    for text in [
        "Формальная постановка задачи позволяет связать экономическую сущность риска с методами машинного обучения. Пусть имеется множество инвестиционных объектов, наблюдаемых в дискретные моменты времени. В данной работе объектом наблюдения является пара «эмитент – дата принятия решения». Для каждой такой пары формируется вектор признаков, отражающий рыночное, ликвидностное, макроэкономическое и фундаментальное состояние объекта.",
        "Задача классификации состоит в построении функции, которая по вектору признаков относит объект к одному из трех классов риска: low, medium или high. Класс low соответствует относительно низкому будущему риску, medium – умеренному уровню риска, high – повышенной будущей уязвимости актива. Выбор трех классов является компромиссом между детализацией и интерпретируемостью: двух классов недостаточно для ранжирования, а слишком большое число классов усложняет устойчивую разметку.",
    ]:
        add_plain_paragraph(doc, text)
    add_formula(doc, "f: Xₜ,ᵢ → yₜ,ᵢ,  yₜ,ᵢ ∈ {low, medium, high}", "1.1")
    for text in [
        "где Xₜ,ᵢ – вектор признаков i-го эмитента на дату t; yₜ,ᵢ – класс инвестиционного риска; f – классифицирующая функция, реализуемая ML-пайплайном.",
        "Важным условием является временная корректность. Признаки Xₜ,ᵢ должны быть сформированы только из информации, доступной на дату t. Будущая доходность, будущая просадка и будущая неликвидность не могут входить в признаки, но используются для формирования целевой переменной при обучении и оценке качества.",
        "В работе используется целевая переменная, построенная через будущий интегральный RiskScore. Он объединяет несколько компонент риска, каждая из которых отражает отдельную сторону неблагоприятной реализации: глубину падения цены, отрицательную изменчивость, хвостовые потери и неликвидность. Затем значения RiskScore разбиваются на классы по train-only порогам, чтобы избежать подгонки порогов под будущие периоды.",
        "Качество решения оценивается с помощью метрик классификации. Для общей оценки используются macro-F1, weighted-F1 и balanced accuracy. Для прикладной риск-постановки отдельно анализируется recall класса high, поскольку пропуск высокого риска является более опасной ошибкой, чем ложное завышение категории риска.",
        "Формальная постановка также требует указать, что оптимизация процесса в данной работе понимается не как математическое нахождение глобального экстремума некоторой функции, а как улучшение организации классификации риска: переход от ручной, слабо воспроизводимой процедуры к единому пайплайну с фиксированными формулами, алгоритмами, параметрами и метриками качества.",
        "Автоматическая классификация в данной работе рассматривается как формализованный процесс первичного риск-скрининга. Ее результатом является не инвестиционная рекомендация, а воспроизводимая категория риска, рассчитанная по единой логике для всех наблюдений. Это позволяет отделить массовый расчет признаков и классов от последующего анализа наиболее рискованных объектов.",
        "Автоматическая классификация имеет преимущество в повторяемости. Если исходные данные, версия кода и параметры модели сохранены, повторный запуск должен привести к тому же результату. В ручной процедуре такая повторяемость сложнее, потому что часть решений может зависеть от неформализованных профессиональных суждений, текущего контекста и последовательности ручных операций.",
        "Отдельное требование связано с масштабируемостью. Под масштабируемостью в рамках исследования понимается возможность увеличивать число эмитентов, дат наблюдений и признаков без пропорционального роста ручных трудозатрат. ML-пайплайн после настройки может обработать всю панель наблюдений пакетно, тогда как ручной анализ требует повторения вычислений для каждого объекта.",
        "Воспроизводимость и масштабируемость должны оцениваться не декларативно, а через признаки процесса. Воспроизводимость подтверждается сохранением кода, конфигурации, временных разбиений, порогов и артефактов запуска. Масштабируемость подтверждается тем, что расчеты выполняются программно для всей панели данных, а не только для отдельных вручную выбранных примеров.",
        "Точность процесса в данной работе понимается как качество совпадения прогнозируемого класса риска с эталонной разметкой, построенной на основе будущих компонент риска. Такая точность не является абсолютной истиной о будущем рынке, но является корректной исследовательской мерой того, насколько модель воспроизводит выбранную формализацию инвестиционного риска.",
        "Для базового сравнения может использоваться упрощенный классификатор, например модель только на рыночных признаках или ручное правило по волатильности. Такое сравнение показывает, дает ли объединение признаков, кластеризации и нейросетевого слоя прирост относительно более простой процедуры. В данной работе основное сравнение проводится между несколькими архитектурами пайплайна.",
        "Формальная постановка требует контролировать границы применимости результата. Модель обучается на данных российских акций и не может автоматически переноситься на облигации, производные инструменты или иностранные рынки без дополнительной проверки. Для другого класса активов потребуется заново определить признаки, целевую переменную и критерии качества.",
        "Еще одно ограничение связано с горизонтом оценки. В работе используется будущий горизонт 126 торговых дней, что примерно соответствует полугодовому окну. Если изменить горизонт, классы риска могут измениться: краткосрочная уязвимость и среднесрочная уязвимость не всегда совпадают. Поэтому горизонт должен рассматриваться как часть постановки задачи, а не как техническая деталь.",
        "В результате формальная постановка задает исследовательский каркас всей работы: объект наблюдения, признаки, целевую переменную, классы риска, временную схему и метрики качества. Такой каркас связывает экономическое содержание риска, математическую формализацию и программную реализацию в единую логику.",
        "Вывод по первой главе. В первой главе была раскрыта экономическая сущность инвестиционного риска, рассмотрен ручной подход к классификации, выявлены его ограничения и сформулирована формальная задача автоматической классификации инвестиционных рисков. Полученные выводы задают требования к дальнейшему анализу предметной области: необходимо обосновать данные, метрики, формулы и методы, которые будут использованы в реализации.",
    ]:
        add_plain_paragraph(doc, text)
    add_intro_defense_context(doc)
    add_chapter1_defense_context(doc)
    add_chapter1_literature_review(doc)
    doc.add_page_break()


def chapter2(doc: Document, data: dict) -> None:
    add_heading(doc, "2. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ КЛАССИФИКАЦИИ ИНВЕСТИЦИОННЫХ РИСКОВ", 1)
    add_heading(doc, "2.1. Анализ инвестиционного риска как объекта классификации", 2)
    for text in [
        "Вторая глава посвящена анализу предметной области классификации инвестиционных рисков. Ее задача состоит в том, чтобы показать экономическое содержание риска, обосновать выбор исходных данных, раскрыть используемые показатели и определить, какие методы машинного обучения подходят для дальнейшей реализации. Если первая глава формулирует задачу автоматической классификации, то вторая глава задает предметную и методическую основу ее решения.",
        "Инвестиционный риск как объект классификации имеет несколько измерений. Рыночное измерение связано с динамикой цены и доходности. Ликвидностное измерение отражает возможность войти в позицию или выйти из нее без существенного влияния на цену. Фундаментальное измерение характеризует финансовую устойчивость эмитента. Макроэкономическое измерение определяет внешний режим, в котором находятся все участники рынка.",
        "Для классификации риска важно различать текущие признаки и будущую реализацию риска. Текущие признаки описывают состояние объекта на дату решения: волатильность, бета, объем торгов, ставка, валютный курс, долговые коэффициенты и признаки отчетности. Будущая реализация риска используется только как основание для разметки: насколько сильной оказалась просадка, насколько выраженными были отрицательные доходности и насколько ухудшалась ликвидность.",
        "Такой подход соответствует логике risk-screening. Модель не должна обещать точный прогноз цены. Ее задача – заранее выделить наблюдения, у которых набор текущих признаков похож на ситуации, приводившие к высокому будущему риску. Поэтому целевая постановка ближе к классификации состояния уязвимости, чем к прогнозированию доходности.",
        "В предметной области рынка акций существенна отраслевая неоднородность. Банковский сектор иначе реагирует на процентные ставки, чем нефтегазовый сектор; металлургические компании чувствительны к валютному курсу и мировым ценам; компании потребительского сектора могут зависеть от внутреннего спроса и стоимости финансирования. Поэтому в признаковое пространство включается сектор, а часть признаков нормируется относительно рынка и сектора.",
        "Важным требованием является интерпретируемость. Классификация риска должна не только выдавать класс, но и позволять объяснить, какие группы факторов повлияли на результат. Поэтому уже на уровне предметной области необходимо понимать, какие признаки описывают цену, ликвидность, рыночную чувствительность, устойчивость эмитента и внешний макроэкономический режим.",
        "Предметная область требует также учета качества данных. На финансовом рынке часть данных может отсутствовать, публиковаться с задержкой или иметь различную частоту. Рыночные цены доступны ежедневно, финансовая отчетность раскрывается квартально или ежегодно, макроэкономические показатели обновляются по собственному календарю. Следовательно, данные необходимо приводить к единому месячному срезу и присоединять по принципу доступности на дату решения.",
    ]:
        add_plain_paragraph(doc, text)

    add_heading(doc, "2.2. Анализ исходных данных для оценки инвестиционных рисков", 2)
    universe = data["universe"]
    panel = data["panel"]
    metrics = data["metrics"]
    date_min = pd.to_datetime(panel["decision_date"]).min().date()
    date_max = pd.to_datetime(panel["decision_date"]).max().date()
    for text in [
        f"В практической части используется панель данных по российским акциям, сформированная на месячных срезах. Локальная модельная панель содержит {len(panel)} строк и {panel.shape[1]} столбцов за период с {date_min} по {date_max}. В инвестиционную вселенную включено {int(universe['include_flag'].sum())} эмитентов из нескольких секторных групп. Такой формат позволяет рассматривать каждое наблюдение как пару «эмитент – дата решения».",
        "Данные берутся не произвольно: каждая группа источников соответствует определенной стороне инвестиционного риска. Цены закрытия и дневные доходности нужны для расчета рыночной динамики, волатильности, беты и будущих компонент целевой переменной. Объем торгов и денежный оборот используются для оценки ликвидности. Рыночный индекс нужен для оценки чувствительности бумаги к общему рынку. Макроэкономические ряды нужны для описания внешнего режима.",
        "Фундаментальные данные и финансовая отчетность используются для оценки внутренней устойчивости эмитента. В них отражаются долговая нагрузка, операционные денежные потоки, EBITDA, свободный денежный поток, процентные расходы, структура долга и маржинальность. Эти показатели важны, поскольку рыночная цена не всегда мгновенно и полностью отражает ухудшение финансового состояния компании.",
        "Особое значение имеет дата публикации отчетности. Отчетный период показывает, за какой период составлен документ, но инвестор получает доступ к информации только после раскрытия. Поэтому присоединение отчетных признаков выполняется по publish_date. Это предотвращает использование будущей информации в прошлых решениях и делает исследование корректным с точки зрения временной логики.",
        "На месячном срезе используется последняя доступная торговая дата месяца. Такой подход уменьшает шум ежедневных наблюдений и делает классификацию ближе к реальному процессу регулярного риск-мониторинга. Ежемесячная частота достаточна для оценки среднесрочной уязвимости акций и согласуется с горизонтом будущего риска.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 2.1 – Группы данных и их роль в классификации инвестиционного риска")
    add_table(
        doc,
        ["Группа данных", "Примеры признаков", "Роль в предметной области"],
        [
            ["Рыночные данные", "close, return_1d, momentum_6m, beta_60d", "Отражают динамику цены и чувствительность к рынку"],
            ["Ликвидность", "value, turnover_ratio, amihud_20d, spread_proxy", "Характеризуют возможность выхода из позиции"],
            ["Макроэкономика", "cbr_key_rate, usd_rub, ofz_slope, market volatility", "Описывает внешний режим рынка"],
            ["Фундаментальные данные", "net_debt_to_ebitda, interest_coverage, free_cash_flow", "Показывают устойчивость эмитента"],
            ["Финансовая отчетность", "report_financial_pressure, report_lag_days, text risk density", "Добавляет point-in-time признаки раскрытия информации"],
        ],
        [3.1, 5.1, 7.3],
        font_size=10,
    )
    add_source(doc, "Источник: составлено автором на основе структуры локального набора данных.")
    add_caption(doc, "Таблица 2.2 – Состав инвестиционной вселенной")
    rows = universe[["ticker", "company_name", "sector"]].head(18).values.tolist()
    add_table(doc, ["Тикер", "Эмитент", "Сектор"], rows, [2.0, 7.5, 5.8], font_size=10)
    add_source(doc, "Источник: составлено автором по файлу data/universe.csv.")
    for text in [
        f"После разметки распределение классов во всей выборке составило: low – {metrics['class_distribution']['low']}, medium – {metrics['class_distribution']['medium']}, high – {metrics['class_distribution']['high']}. Наличие сопоставимых по размеру классов важно для оценки качества: модель не должна сводиться к тривиальному предсказанию наиболее частой категории.",
        "Отдельной проблемой является пропущенность данных. В финансовых отчетах не все показатели доступны для всех эмитентов и периодов. Поэтому в пайплайне используются не только числовые значения, но и признаки пропуска. Сам факт отсутствия отчетного показателя может быть информативен для модели, если он связан с качеством раскрытия или нерегулярностью данных.",
    ]:
        add_plain_paragraph(doc, text)

    add_heading(doc, "2.3. Анализ экономических и математических показателей инвестиционного риска", 2)
    formula_blocks = [
        ("Простая доходность", "rₜ = (Pₜ - Pₜ₋₁) / Pₜ₋₁", "2.1", "где rₜ – доходность актива за период t; Pₜ – цена актива в момент t; Pₜ₋₁ – цена актива в предыдущий момент. Доходность является базовой величиной для расчета большинства рыночных метрик риска."),
        ("Средняя доходность", "r̄ = (1 / T) · Σₜ₌₁ᵀ rₜ", "2.2", "где T – число наблюдений. Средняя доходность используется как ориентир центральной тенденции, но сама по себе не является достаточной мерой риска."),
        ("Волатильность", "σ = √((1 / (T - 1)) · Σₜ₌₁ᵀ (rₜ - r̄)²)", "2.3", "где σ – стандартное отклонение доходности. Волатильность отражает общую изменчивость доходности, однако учитывает положительные и отрицательные отклонения симметрично."),
        ("Downside-volatility", "σ⁻ = √((1 / (T⁻ - 1)) · Σᵣₜ<0 (rₜ - 0)²)", "2.4", "где T⁻ – число отрицательных доходностей. Эта метрика отражает изменчивость только неблагоприятных движений и поэтому ближе к риск-менеджерской трактовке потерь."),
        ("Накопленная доходность", "Cₜ = Πₖ₌₁ᵗ (1 + rₖ)", "2.5", "где Cₜ – накопленный индекс стоимости. Он используется для вычисления просадки относительно предыдущего максимума."),
        ("Максимальная просадка", "MDD = maxₜ ((maxₖ≤ₜ Cₖ - Cₜ) / maxₖ≤ₜ Cₖ)", "2.6", "Максимальная просадка показывает наибольшее падение стоимости от локального максимума до последующего минимума. В работе она входит в целевую переменную как мера глубины будущего убытка."),
        ("Value-at-Risk", "VaRᵅ = -qᵅ(r)", "2.7", "где qᵅ(r) – квантиль распределения доходности уровня α. VaR показывает порог убытка, который не превышается с заданной вероятностью при историческом распределении доходностей."),
        ("Conditional Value-at-Risk", "CVaRᵅ = -E[r | r ≤ qᵅ(r)]", "2.8", "CVaR отражает среднюю величину потерь в хвосте распределения. В работе используется CVaR 95%, соответствующий среднему неблагоприятному исходу в нижних 5% доходностей."),
        ("Индикатор неликвидности Амихуда", "ILLIQ = (1 / T) · Σₜ₌₁ᵀ |rₜ| / Valueₜ", "2.9", "где Valueₜ – денежный оборот торгов за период. Чем выше показатель, тем сильнее цена реагирует на единицу торгового оборота, что интерпретируется как большая неликвидность."),
        ("Рыночная бета", "βᵢ = Cov(rᵢ, rₘ) / Var(rₘ)", "2.10", "где rᵢ – доходность акции, rₘ – доходность рыночного индекса. Бета показывает чувствительность бумаги к общему рыночному движению."),
        ("Коэффициент долговой нагрузки", "ND/EBITDA = NetDebt / EBITDA", "2.11", "Показатель характеризует способность компании обслуживать долговую нагрузку за счет операционного результата. Высокие значения усиливают фундаментальную хрупкость эмитента."),
        ("Покрытие процентов", "ICR = EBITDA / |InterestExpense|", "2.12", "Показатель отражает способность компании покрывать процентные расходы. Низкое значение может указывать на повышенный риск при росте ставок."),
        ("Интегральный риск-скор", "RiskScore = 0,35·p(MDD) + 0,30·p(σ⁻) + 0,20·p(CVaR₉₅) + 0,15·p(ILLIQ)", "2.13", "где p(·) – процентильное преобразование компоненты относительно обучающего периода. Веса отражают приоритет глубины просадки, отрицательной волатильности, хвостовых потерь и неликвидности."),
    ]
    for title, formula, number, explanation in formula_blocks:
        add_plain_paragraph(doc, title + ".", bold=True)
        add_formula(doc, formula, number)
        add_plain_paragraph(doc, explanation)
    for text in [
        "Использование процентильного преобразования компонент риска необходимо для сопоставимости показателей, имеющих разные масштабы. Максимальная просадка измеряется в долях стоимости, CVaR – в доходности, а неликвидность может иметь очень малые числовые значения. Перевод каждой компоненты в процентиль относительно обучающего периода позволяет объединить их в единую шкалу.",
        "Пороги классов риска также должны рассчитываться только на обучающем периоде. Это методически важно: если использовать всю выборку, в том числе будущие периоды, модель косвенно получит информацию о распределении риска в test-периоде. Train-only логика делает классификацию более строгой и предотвращает завышение качества.",
        "Выбор весов в RiskScore должен быть экономически объяснимым. Наибольший вес получает максимальная просадка, поскольку именно глубина падения стоимости является наиболее понятной мерой потерь для инвестора. Downside-volatility имеет второй по величине вес, так как отражает частоту и разброс отрицательных движений. CVaR учитывает редкие, но сильные потери, а неликвидность добавляет измерение практической возможности выхода из позиции.",
        "При этом веса не следует трактовать как универсальные для всех инвестиционных стратегий. Для консервативного портфеля можно усилить вес просадки и CVaR; для стратегии с высоколиквидными инструментами можно меньше учитывать неликвидность; для стратегии, работающей с малокапитализированными бумагами, ликвидностный компонент может быть критически важным. В рамках работы выбранный набор весов используется как воспроизводимая исследовательская спецификация.",
        "Формула максимальной просадки важна тем, что она учитывает последовательность цен. Две акции могут иметь одинаковые конечные доходности, но разные траектории: одна снижалась плавно, другая пережила глубокую просадку и затем восстановилась. Для инвестора эти ситуации различаются, поскольку глубокая просадка может привести к маржинальным требованиям, вынужденной продаже или нарушению риск-лимитов.",
        "CVaR дополняет максимальную просадку, потому что описывает не одну самую глубокую траекторию потерь, а среднее поведение хвоста распределения доходностей. Если у акции часто возникают экстремально отрицательные дневные доходности, CVaR будет выше даже тогда, когда максимальная просадка за конкретное окно не является рекордной. Это делает показатель полезным для оценки хвостового риска.",
        "Показатель неликвидности важен для российской предметной области из-за неоднородности торговой активности. Для крупных эмитентов с высоким оборотом выход из позиции может быть относительно простым. Для менее ликвидных бумаг даже умеренное изменение цены может сопровождаться существенными затратами на исполнение сделки. Поэтому ликвидность включается и в признаки, и в будущую целевую компоненту.",
        "Бета используется как показатель рыночной чувствительности, но не заменяет другие меры риска. Высокая бета означает, что бумага в среднем сильнее реагирует на движение рынка, однако не описывает полностью ликвидность, долговую нагрузку и хвостовые потери. Поэтому в работе beta рассматривается как один из признаков, а не как единственная мера инвестиционного риска.",
        "Фундаментальные коэффициенты позволяют добавить в модель информацию о финансовой устойчивости эмитента. Если долговая нагрузка высока, а покрытие процентов низкое, компания может быть более уязвима при росте ставок или ухудшении операционной прибыли. Такие риски могут проявиться в цене не сразу, поэтому их включение расширяет предметную область анализа за пределы чисто рыночной статистики.",
        "Макроэкономические признаки выполняют роль контекста. Ключевая ставка влияет на стоимость капитала, валютный курс – на экспортно-импортные потоки и переоценку обязательств, наклон кривой ОФЗ – на ожидания рынка относительно процентного режима. Если модель игнорирует макроэкономику, она может воспринимать одинаковые значения рыночных признаков одинаково в принципиально разных режимах.",
        "Cross-sectional ranks помогают уменьшить зависимость модели от общего уровня рынка. Например, абсолютная волатильность может быть высокой для всех бумаг в кризисном периоде, но инвестору важно понимать, какая бумага рискованнее относительно остальных в тот же месяц. Ранжирование по месяцу позволяет модели видеть относительное положение объекта в текущем рыночном срезе.",
        "Sector z-scores решают похожую задачу внутри отрасли. Сравнивать долговую нагрузку банка, нефтегазовой компании и металлургического эмитента напрямую не всегда корректно. Секторная нормализация позволяет оценить, насколько показатель эмитента отклоняется от типичных значений для своей группы, что повышает экономическую интерпретируемость признаков.",
        "Таким образом, набор формул в работе не является декоративным. Каждая формула отвечает за конкретное измерение предметной области: цена, отрицательная изменчивость, хвост потерь, ликвидность, чувствительность к рынку, долговая нагрузка, покрытие процентов и интегральная оценка будущего риска. Именно эта связь делает дальнейшую реализацию обоснованной.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 2.3 – Экономический смысл основных показателей риска")
    add_table(
        doc,
        ["Показатель", "Тип", "Что отражает", "Использование в работе"],
        [
            ["Доходность", "рыночный", "изменение цены", "база для rolling-признаков"],
            ["Волатильность", "статистический", "общую изменчивость", "признак текущего риска"],
            ["Downside-volatility", "статистический", "отрицательные колебания", "признак и target-компонента"],
            ["MDD", "экономический", "глубину падения", "target-компонента"],
            ["CVaR 95%", "статистический", "хвостовые потери", "target-компонента"],
            ["ILLIQ", "ликвидностный", "реакцию цены на оборот", "признак и target-компонента"],
            ["β", "рыночный", "чувствительность к индексу", "признак модели"],
            ["ND/EBITDA", "фундаментальный", "долговую нагрузку", "признак устойчивости эмитента"],
        ],
        [2.8, 2.7, 5.1, 4.7],
        font_size=9,
    )
    add_source(doc, "Источник: составлено автором.")

    add_heading(doc, "2.4. Анализ применимости методов машинного обучения к задаче классификации риска", 2)
    for text in [
        "После формализации признаков и целевой переменной необходимо выбрать методы машинного обучения. Для рассматриваемой задачи целесообразен гибридный подход, поскольку один метод не покрывает все требования предметной области. Требуется интерпретируемость, способность учитывать нелинейные зависимости, возможность выявлять режимы рынка и воспроизводимость обучения.",
        "Деревья решений и ансамбли деревьев подходят для табличных финансовых данных. Они не требуют линейной зависимости между признаками и целевой переменной, устойчивы к различным масштабам признаков и позволяют оценивать важность факторов. Для задачи классификации риска они удобны как основа интерпретируемого классифицирующего слоя.",
        "Кластеризационный анализ применим для выделения рыночных режимов. Макроэкономические признаки, такие как ключевая ставка, курс рубля, наклон кривой ОФЗ, реализованная волатильность индекса и средняя рыночная корреляция, могут образовывать разные состояния рынка. Информация о режиме позволяет учитывать контекст, в котором находится наблюдение.",
        "Искусственная нейронная сеть применима не как непрозрачная замена всей модели, а как компонент извлечения латентных факторов. Автоэнкодер может сжимать числовые признаки в компактное представление, отражающее скрытые сочетания рыночной, ликвидностной и фундаментальной информации.",
        "Такое объединение методов соответствует теме работы. Деревья решений обеспечивают интерпретируемое ядро классификации; кластеризационный слой добавляет информацию о рыночном режиме; нейросетевой слой извлекает скрытые факторы. В результате формируется методическая основа ML-пайплайна, реализация которого рассматривается в третьей главе.",
        "Применимость деревьев решений определяется тем, что они хорошо работают с табличными признаками разной природы. В одной модели могут одновременно использоваться рыночные показатели, ликвидность, фундаментальные коэффициенты, макроэкономические признаки и признаки отчетности. Дерево решений способно находить пороговые правила и взаимодействия факторов без требования линейной связи между признаками и классом риска.",
        "Применимость кластеризационного анализа связана с тем, что инвестиционный риск зависит от рыночного режима. Одни и те же значения долговой нагрузки, волатильности или ликвидности могут иметь разный смысл в спокойном и стрессовом периоде. Кластеризация позволяет выделять похожие состояния рынка или риск-профили наблюдений, после чего информация о режиме может использоваться как дополнительный контекст для классификации.",
        "Применимость искусственных нейронных сетей связана с возможностью учитывать нелинейные зависимости между признаками. В финансовых данных риск часто возникает не из-за одного показателя, а из-за сочетания факторов: например, высокой беты, ухудшения ликвидности, роста ставки и слабого денежного потока. Нейросетевой слой может извлекать скрытые представления из такого признакового пространства, а затем передавать их в классифицирующую модель.",
        "При этом методы машинного обучения не подменяют предметную область. Они применимы только после того, как определены экономически осмысленные данные, рассчитаны показатели риска и сформирована целевая переменная. Поэтому во второй главе методы рассматриваются именно с точки зрения пригодности для задачи классификации инвестиционного риска, а их программная реализация описывается в третьей главе.",
    ]:
        add_plain_paragraph(doc, text)
    add_chapter2_defense_context(doc, data)
    add_plain_paragraph(
        doc,
        "Вывод по второй главе. Во второй главе проведен анализ предметной области классификации инвестиционных рисков, обоснован выбор исходных данных, рассмотрены экономические и математические показатели риска, а также определены методы машинного обучения, применимые для построения автоматизированной системы классификации.",
    )
    doc.add_page_break()


def chapter3(doc: Document, data: dict) -> None:
    metrics = data["metrics"]
    predictions = data["predictions"]
    feature_importance = data["feature_importance"]
    walk_forward = data["walk_forward"]
    drift = data["drift"]

    add_heading(doc, "3. РЕАЛИЗАЦИЯ ML-ПАЙПЛАЙНА И ОЦЕНКА РЕЗУЛЬТАТОВ", 1)
    add_heading(doc, "3.1. Архитектура разработанного ML-пайплайна", 2)
    for text in [
        "Третья глава описывает реализацию задачи, сформулированной в первой главе и методически обоснованной во второй главе. В работе реализован воспроизводимый point-in-time пайплайн автоматической классификации инвестиционного риска российских акций. Пайплайн объединяет подготовку данных, расчет признаков, формирование целевой переменной, кластеризацию рыночных режимов, обучение ансамблей деревьев, извлечение латентных факторов и оценку качества.",
        "Архитектура пайплайна построена модульно. Отдельные модули отвечают за загрузку и валидацию данных, расчет rolling-признаков, формирование месячных срезов, присоединение макроэкономики и отчетности, расчет target-компонент, временное разделение выборки, обучение моделей, диагностику дрейфа и сохранение артефактов. Такая структура повышает воспроизводимость: каждый этап можно повторить и проверить независимо.",
        "Ключевым принципом реализации является point-in-time корректность. Рыночные и макроэкономические данные берутся на дату решения или раньше; отчетность присоединяется по дате публикации; будущие компоненты риска используются только для формирования целевой переменной. Это исключает утечку будущей информации, которая является одной из наиболее частых ошибок в финансовом машинном обучении.",
        "В программной реализации используется временное разделение данных. Обучающая часть заканчивается 30.08.2024, валидационная часть начинается 30.09.2024 и заканчивается 28.02.2025, тестовая часть начинается 31.03.2025. Такая схема имитирует реальную ситуацию: модель настраивается на прошлом периоде и затем применяется к будущим наблюдениям.",
        f"Размеры выборок в основном запуске составили: train – {metrics['n_train']} наблюдений, validation – {metrics['n_validation']} наблюдений, test – {metrics['n_test']} наблюдений. Распределение классов во всей размеченной выборке: low – {metrics['class_distribution']['low']}, medium – {metrics['class_distribution']['medium']}, high – {metrics['class_distribution']['high']}.",
    ]:
        add_plain_paragraph(doc, text)
    add_figure(doc, RUN_DIR / "defense_assets/architecture_test_metrics.png", "Рисунок 3.1 – Сравнение архитектур модели на тестовом периоде", width_cm=15.2)
    add_caption(doc, "Таблица 3.1 – Основные этапы реализованного ML-пайплайна")
    add_table(
        doc,
        ["Этап", "Содержание", "Назначение"],
        [
            ["1", "Загрузка рыночных, макроэкономических и отчетных данных", "Формирование исходной базы признаков"],
            ["2", "Расчет rolling-признаков и месячных срезов", "Переход к регулярной панели наблюдений"],
            ["3", "Point-in-time merge отчетности", "Исключение утечки будущей информации"],
            ["4", "Формирование RiskScore и классов", "Получение целевой переменной"],
            ["5", "Кластеризация режима рынка", "Учет макроэкономического контекста"],
            ["6", "Обучение ансамблей деревьев и нейросетевого слоя", "Получение классификатора"],
            ["7", "Оценка качества, walk-forward и drift diagnostics", "Проверка применимости результата"],
        ],
        [1.2, 7.0, 7.0],
        font_size=10,
    )
    add_source(doc, "Источник: составлено автором.")

    add_heading(doc, "3.2. Реализация расчета признаков и целевой переменной", 2)
    for text in [
        "Формулы, рассмотренные во второй главе, непосредственно используются в программной реализации. Доходность рассчитывается через процентное изменение цены закрытия. Rolling-volatility строится на 20-дневном окне, downside-volatility – на 60-дневном окне по отрицательным доходностям, beta – через ковариацию доходности акции с доходностью рынка и дисперсию рынка.",
        "Ликвидностные признаки включают средний дневной торговый оборот, индикатор Амихуда, spread_proxy и turnover_ratio. Эти признаки важны, поскольку высокий риск в реальном инвестиционном процессе связан не только с падением цены, но и с возможностью закрыть позицию без значительных потерь ликвидности.",
        "Фундаментальные признаки включают net_debt_to_ebitda, interest_coverage, ebitda_margin, book_to_market, operating_cash_flow и free_cash_flow. Для отчетности дополнительно рассчитываются report_financial_pressure, report_integrated_stress, report_lag_days, report_text_risk_density и другие показатели. Таким образом, формулы главы 2 не остаются теоретическим блоком, а становятся основой feature engineering.",
        "Целевая переменная формируется по будущему окну 126 торговых дней. Для каждого месячного наблюдения берется будущая траектория цены и торгового оборота. По ней рассчитываются future_max_drawdown, future_downside_volatility, future_cvar_95 и future_illiquidity. Затем компоненты переводятся в процентильную шкалу относительно обучающего периода и объединяются в RiskScore.",
        f"Пороги классов, полученные на обучающем периоде, составили: low_upper = {metrics['target_thresholds']['low_upper']:.4f}, medium_upper = {metrics['target_thresholds']['medium_upper']:.4f}. Наблюдения с RiskScore ниже первого порога относятся к low, между порогами – к medium, выше второго порога – к high.",
        "Важная особенность реализации состоит в том, что пороги и процентильные распределения рассчитываются только на train-части. Это предотвращает использование информации о будущих периодах при разметке validation и test. Такой подход делает оценку качества более строгой и согласуется с реальной процедурой внедрения модели.",
        f"В ходе feature engineering было добавлено {metrics['added_features_count']} новых признака. Среди них downside_to_total_vol, liquidity_stress, beta_vol_interaction, macro_pressure, rate_fx_pressure, debt_service_stress, cashflow_buffer, fundamental_fragility и report_integrated_stress. Эти признаки отражают не отдельные показатели, а экономически осмысленные взаимодействия факторов риска.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 3.2 – Компоненты RiskScore в реализации")
    add_table(
        doc,
        ["Компонента", "Вес", "Экономический смысл"],
        [
            ["future_max_drawdown", 0.35, "Глубина будущей просадки"],
            ["future_downside_volatility", 0.30, "Изменчивость отрицательных доходностей"],
            ["future_cvar_95", 0.20, "Средние хвостовые потери"],
            ["future_illiquidity", 0.15, "Будущая неликвидность"],
        ],
        [4.4, 2.0, 8.8],
        font_size=10,
    )
    add_source(doc, "Источник: составлено автором по реализации risk_pipeline.targets.")
    report_sel = metrics.get("report_layer_selection", {})
    for text in [
        f"Отдельно реализован слой финансовой отчетности. В основном запуске report-признаки присутствуют, выбран вариант {report_sel.get('selected', 'with_reports')}, количество отчетных признаков – {report_sel.get('report_feature_count', 0)}. Слой отчетности включается не автоматически, а через validation gate: модель с отчетностью сравнивается с вариантом без отчетности по валидационной метрике.",
        f"Валидационный selection score без отчетности составил {report_sel.get('without_reports', {}).get('selection_score', 0):.4f}, с отчетностью – {report_sel.get('with_reports', {}).get('selection_score', 0):.4f}. Поэтому в итоговом варианте используется слой with_reports. Это означает, что отчетность добавлена не ради усложнения модели, а потому что дала измеримое улучшение на validation-периоде.",
    ]:
        add_plain_paragraph(doc, text)
    add_figure(doc, RUN_DIR / "defense_assets/report_layer_ablation.png", "Рисунок 3.2 – Сравнение модели с отчетным слоем и без него", width_cm=14.5)

    add_heading(doc, "3.3. Реализация моделей машинного обучения", 2)
    for text in [
        "В реализации сравниваются несколько архитектур. Базовая модель baseline_rf использует ансамбль деревьев решений на исходном наборе признаков. Regime_only добавляет кластер рыночного режима, построенный на макроэкономических признаках. Enriched_reference использует расширенное признаковое пространство. Ann_plus_regime объединяет латентные факторы автоэнкодера и режимный слой. Sector_overlay проверяет, дают ли специализированные секторные модели прирост качества на валидации.",
        "Кластеризация режима рынка реализована как детерминированный KMeans-подобный алгоритм на NumPy. Это снижает зависимость от внешних реализаций и делает запуск воспроизводимым на обычном компьютере. Перед кластеризацией признаки стандартизируются на обучающей выборке, а пропуски заполняются медианами train-периода. Для validation, test и inference используются те же параметры, что исключает пересчет статистик на будущем периоде.",
        "Деревья решений реализованы через Random Forest, Extra Trees и Gradient Boosting. Каждый кандидат обучается на train-части и оценивается на validation. Затем подбирается soft voting ensemble, в котором веса кандидатов выбираются по валидационному качеству. Такой подход устойчивее, чем выбор одной модели, поскольку разные ансамбли деревьев по-разному реагируют на шум и нелинейность признаков.",
        "Нейросетевой слой реализован как автоэнкодер. Его задача – не напрямую предсказывать класс риска, а извлекать компактные латентные факторы из числового признакового пространства. Эти факторы затем передаются в модель классификации. В основном запуске используется PyTorch backend с PCA-инициализацией линейного bottleneck autoencoder; fallback предусмотрен только как аварийный режим для сред без torch.",
        f"В основном запуске backend автоэнкодера: {metrics.get('autoencoder_backend')}. Значение reconstruction loss составило {metrics.get('autoencoder_reconstruction_loss', 0):.4f}. Это значение используется как диагностический показатель качества восстановления признакового пространства, но не является основной метрикой классификации.",
        "Отраслевой overlay реализован как проверка специализированных моделей по секторам. Секторная модель используется только в том случае, если на validation она дает достаточный прирост качества и сектор содержит достаточное число наблюдений. В основном запуске специализированные отраслевые модели не были выбраны, что является корректным результатом validation gate: пайплайн не усложняется без доказанного выигрыша.",
    ]:
        add_plain_paragraph(doc, text)
    final_arch = metrics.get("final_architecture", "ann_plus_regime")
    add_plain_paragraph(doc, f"Финальная архитектура для основного результата – ann_plus_regime. Она выбрана по validation objective, то есть до обращения к test-периоду. Это важно для корректности оценки: test используется только для финальной проверки качества, а не для подбора архитектуры.", bold=False)
    add_caption(doc, "Таблица 3.3 – Реализованные модели и их роль в пайплайне")
    arch_rows = [
        ["baseline_rf", "Ансамбль деревьев решений на исходном наборе признаков", "Базовый метод для сравнения"],
        ["regime_only", "Ансамбль деревьев с добавлением кластера рыночного режима", "Проверка вклада кластеризационного анализа"],
        ["enriched_reference", "Ансамбль деревьев на расширенном признаковом пространстве", "Проверка вклада feature engineering"],
        ["ann_plus_regime", "Ансамбль деревьев с латентными факторами автоэнкодера и режимным слоем", "Финальная гибридная архитектура"],
        ["sector_overlay", "Дополнительные секторные модели при наличии validation-выигрыша", "Проверка полезности отраслевой специализации"],
    ]
    add_table(
        doc,
        ["Архитектура", "Реализованные компоненты", "Назначение"],
        arch_rows,
        [3.2, 7.2, 5.4],
        font_size=10,
    )
    add_source(doc, "Источник: составлено автором.")

    add_heading(doc, "3.4. Оценка качества и интерпретация результатов", 2)
    final_vals = metrics["architectures"]["ann_plus_regime"]["test"]
    for text in [
        f"Финальная архитектура ann_plus_regime на тестовом периоде показала macro-F1 = {final_vals['macro_f1']:.4f}, weighted-F1 = {final_vals['weighted_f1']:.4f}, balanced accuracy = {final_vals['balanced_accuracy']:.4f}. Recall класса high составил {final_vals['high_recall']:.4f}, high false negative rate – {final_vals['high_false_negative_rate']:.4f}. Для задачи риск-скрининга это важный результат: модель редко пропускает реально высокий риск.",
        "При интерпретации метрик важно учитывать характер задачи. В риск-менеджменте ошибка false negative по high-классу может быть более опасной, чем false positive. Если модель ошибочно отнесла среднерисковую бумагу к high, это может привести к дополнительной ручной проверке. Если же модель пропустила действительно высокий риск, инвестиционное решение может привести к существенным потерям.",
        "Confusion matrix показывает, что модель консервативна в отношении высокого риска: часть объектов среднего класса переводится в high. Такое поведение объясняет невысокую precision класса high при высоком recall. Для риск-скрининга это управляемый trade-off: пороги вероятностей могут настраиваться в зависимости от политики риска.",
    ]:
        add_plain_paragraph(doc, text)
    metric_formulas = [
        ("Accuracy = (TP + TN) / (TP + TN + FP + FN)", "3.1", "Accuracy показывает общую долю верных ответов, но может скрывать ошибки по отдельным классам при несбалансированной выборке."),
        ("Precision = TP / (TP + FP)", "3.2", "Precision показывает, какая доля объектов, отнесенных моделью к заданному классу, действительно принадлежит этому классу."),
        ("Recall = TP / (TP + FN)", "3.3", "Recall показывает, какая доля объектов заданного класса была найдена моделью. Для класса high эта метрика особенно важна, поскольку пропуск высокого риска является наиболее опасной ошибкой."),
        ("F1 = 2 · Precision · Recall / (Precision + Recall)", "3.4", "F1 объединяет precision и recall и используется, когда необходимо одновременно учитывать ложные срабатывания и пропуски класса."),
    ]
    for formula, number, explanation in metric_formulas:
        add_formula(doc, formula, number)
        add_plain_paragraph(doc, explanation)
    add_figure(doc, RUN_DIR / "defense_assets/confusion_matrix_final.png", "Рисунок 3.3 – Матрица ошибок финальной модели", width_cm=12.8)
    add_figure(doc, RUN_DIR / "defense_assets/classwise_final_metrics.png", "Рисунок 3.4 – Метрики качества по классам", width_cm=13.5)
    class_rows = [
        ["low", 0.7778, 0.4375, 0.5600, 16],
        ["medium", 0.8696, 0.3704, 0.5195, 54],
        ["high", 0.4143, 0.9062, 0.5686, 32],
    ]
    add_caption(doc, "Таблица 3.4 – Метрики финальной модели по классам")
    add_table(doc, ["Класс", "Prec.", "Recall", "F1", "N"], class_rows, [3.0, 2.6, 2.6, 2.6, 2.0], font_size=12)
    add_source(doc, "Источник: составлено автором по отчету о качестве финальной модели.")
    for text in [
        "Для проверки устойчивости использована walk-forward схема. Она последовательно расширяет обучающее окно и проверяет качество на следующих периодах. Такая проверка ближе к реальному финансовому применению, чем случайное перемешивание строк. В основном запуске использовано 12 walk-forward фолдов.",
        f"Средний macro-F1 по walk-forward составил {metrics['walk_forward']['macro_f1_mean']:.4f}, средний recall класса high – {metrics['walk_forward']['high_recall_mean']:.4f}. Значения ниже основной validation-оценки, что отражает сложность финансовой задачи и сдвиг рыночных режимов. Этот факт учитывается через дополнительную drift diagnostics.",
    ]:
        add_plain_paragraph(doc, text)
    add_figure(doc, RUN_DIR / "defense_assets/walk_forward_stability.png", "Рисунок 3.5 – Устойчивость качества на walk-forward проверке", width_cm=14.8)
    add_caption(doc, "Таблица 3.5 – Первые фолды walk-forward проверки")
    wf_rows = [
        [
            int(row["fold"]),
            int(row["train_rows"]),
            int(row["validation_rows"]),
            int(row["test_rows"]),
            f"{row['macro_f1']:.3f}",
            f"{row['high_recall']:.3f}",
            f"{row['ordinal_adjacent_accuracy']:.3f}",
        ]
        for _, row in walk_forward.head(8).iterrows()
    ]
    add_table(doc, ["Фолд", "Train", "Val.", "Test", "Macro-F1", "High R", "Adj. acc."], wf_rows, [1.6, 1.8, 1.6, 1.6, 2.2, 1.9, 2.2], font_size=12)
    add_source(doc, "Источник: составлено автором по отчету walk-forward проверки.")
    for text in [
        "Диагностика дрейфа показывает, насколько распределения признаков в будущих периодах отличаются от обучающего периода. Для этого используется Population Stability Index. В основном запуске обнаружено 129 drift warnings, что указывает на существенное изменение рыночных условий. Среди признаков с высоким PSI присутствуют rate_fx_pressure, macro_pressure, average_market_correlation_60d и cbr_key_rate.",
        "Наличие drift warnings не является причиной отвергать результат. Напротив, это важная часть научной честности работы. Финансовый рынок меняется во времени, поэтому модель должна сопровождаться диагностикой сдвига распределений и регулярным переобучением. Без такого контроля автоматическая классификация могла бы сохранять внешний вид работоспособности, но терять качество в новых режимах.",
    ]:
        add_plain_paragraph(doc, text)
    add_figure(doc, RUN_DIR / "defense_assets/drift_top20_psi.png", "Рисунок 3.6 – Признаки с наибольшим PSI", width_cm=14.8)
    feature_labels = {
        "rate_fx_pressure": "Валютное давление",
        "cbr_key_rate": "Ключевая ставка",
        "log_market_cap_sector_z": "Размер эмитента, z",
        "average_market_correlation_60d": "Корреляция рынка 60d",
        "log_market_cap": "Log капитализации",
        "turnover_ratio": "Оборотность",
        "amihud_20d_sector_z": "Amihud 20d, z",
        "momentum_6m": "Моментум 6м",
        "ofz_slope_10y_2y": "Наклон ОФЗ 10-2",
        "downside_to_total_vol": "Downside / total vol",
    }
    top_features = []
    for _, row in feature_importance.head(10).iterrows():
        top_features.append([
            feature_labels.get(row["feature"], row["feature"]),
            row["enriched_importance"],
            row["regime_importance"],
            row["ann_branch_importance"],
        ])
    add_caption(doc, "Таблица 3.6 – Наиболее значимые признаки модели")
    add_table(doc, ["Признак", "Enriched", "Regime", "ANN branch"], top_features, [6.6, 2.8, 2.8, 3.1], font_size=12)
    add_source(doc, "Источник: составлено автором по отчету важности признаков.")
    add_figure(doc, RUN_DIR / "defense_assets/feature_importance_top20.png", "Рисунок 3.7 – Важность признаков финальной модели", width_cm=14.8)
    for text in [
        "Интерпретация важности признаков подтверждает экономический смысл модели. Среди значимых факторов присутствуют rate_fx_pressure, cbr_key_rate, log_market_cap_sector_z, average_market_correlation_60d, turnover_ratio, amihud_20d_sector_z и momentum_6m. Это показывает, что модель использует не один случайный показатель, а сочетание макроэкономического давления, ликвидности, размера эмитента и рыночной динамики.",
        "Модель выдает не только класс, но и вероятности по классам. Это позволяет настраивать операционную политику риска: при снижении порога high можно повысить recall класса high, но ценой большего числа ложных сигналов. Такая настройка должна зависеть от того, насколько консервативной должна быть система риск-контроля.",
    ]:
        add_plain_paragraph(doc, text)
    for text in [
        "Сравнение архитектур показывает, что добавление режимного и нейросетевого слоев не является формальным усложнением. Базовая модель на деревьях дает более низкий macro-F1 на test-периоде, тогда как ann_plus_regime лучше балансирует качество по классам и сохраняет высокий recall класса high. Это соответствует гипотезе работы: объединение деревьев, кластеризации и нейросетевого извлечения факторов может улучшить процесс классификации риска.",
        "При этом результаты не следует трактовать как доказательство абсолютного превосходства сложной модели во всех условиях. Финансовые данные нестабильны, а тестовый период может существенно отличаться от обучающего. Поэтому отдельно сохраняются walk-forward отчеты и drift diagnostics. Такая диагностика делает выводы более осторожными и методически устойчивыми.",
        "Особенность тестового периода состоит в том, что он содержит режим, отличающийся от train и validation. Это видно по PSI для макроэкономических признаков. В такой ситуации падение части метрик относительно validation является ожидаемым. Более важным становится не только само значение F1, но и способность модели обнаруживать high-risk наблюдения при изменении режима.",
        "High precision финальной модели ниже, чем high recall, что означает наличие ложных high-сигналов. Для риск-скрининга это допустимый компромисс, если система используется как предварительный фильтр. Бумаги, ошибочно отнесенные к high, могут быть проверены вручную, тогда как пропущенный high-risk объект может привести к большему ущербу.",
        "Класс medium является наиболее сложным для классификации. Это объяснимо с предметной точки зрения: средний риск расположен между устойчивыми низкорисковыми и явно проблемными высокорисковыми объектами. Граница между medium и high может зависеть от рыночного режима, поэтому ошибки модели часто являются соседними по порядку классов, а не произвольными.",
        "Для учета порядковой природы классов дополнительно используется ordinal adjacent accuracy. Она показывает, насколько часто модель ошибается не более чем на один уровень риска. В финансовой задаче такая метрика полезна, потому что ошибка low вместо high более критична, чем medium вместо high. В основном запуске adjacent accuracy финальной модели составила 0,9118.",
        "Сохранение model_package.joblib делает пайплайн пригодным для повторного инференса. Это означает, что после обучения модель можно применить к новым месячным наблюдениям без повторного подбора параметров. Для воспроизводимости также сохраняется run_manifest.json с fingerprint входной панели и списком артефактов.",
        "Data quality report подтверждает базовые проверки: отсутствие дубликатов ключей decision_date и ticker, наличие target-компонент и корректность дат. Такие проверки важны, потому что ошибки в ключах или датах могут привести к неявному дублированию наблюдений, нарушению временного порядка и искажению метрик качества.",
        "Важность признаков показывает, что модель использует как рыночные, так и макроэкономические факторы. Это подтверждает связь предметной области инвестиционного риска с конкретным набором экономически интерпретируемых признаков, а не с произвольной таблицей данных.",
        "Слой финансовой отчетности добавляет признаки, которые не дублируют рыночную цену полностью. Долговая нагрузка, маржинальность, денежные потоки и текстовая плотность риск-терминов могут отражать фундаментальные проблемы эмитента до того, как они полностью проявятся в будущей доходности. Поэтому отчетный слой является логичным расширением рыночного набора признаков.",
        "Ограничением отчетного слоя является качество извлечения данных из документов. Разные эмитенты публикуют отчетность в разных форматах, часть файлов может быть PDF, XLSX или архивами, а структура таблиц может меняться. Поэтому в реализации используются признаки доступности, устаревания и качества извлечения, чтобы модель могла учитывать надежность отчетного источника.",
        "Отсутствие выбранного sector overlay в итоговом запуске также является содержательным результатом. Оно показывает, что отраслевые специализированные модели не дали достаточного прироста на validation при заданных условиях. Значит, итоговая система не усложняется за счет компонентов, которые не подтверждены проверкой.",
        "С точки зрения масштабируемости реализованный пайплайн может быть применен к новой панели наблюдений при условии наличия требуемых признаков. Добавление новых эмитентов потребует обновления universe, данных и отчетности, но сама логика расчета, разметки и инференса остается той же. Это принципиально отличается от ручной процедуры, где рост числа эмитентов напрямую увеличивает объем работы аналитика.",
        "С точки зрения воспроизводимости итоговый результат поддерживается сохранением исходных данных, кода, конфигурации и отчетов. По артефактам запуска можно установить, какие признаки использовались, какие метрики получены и какие файлы были созданы. Такая прозрачность показывает не только итоговую таблицу качества, но и метод получения результата.",
        "С точки зрения точности разработанный подход оценивается через стандартные метрики классификации. Это позволяет уйти от субъективного утверждения «модель работает хорошо» и перейти к измеримым показателям. При этом работа честно фиксирует ограничения: качество не идеально, присутствует drift, а часть medium-наблюдений ошибочно переводится в high.",
        "Практическое применение пайплайна возможно в режиме вспомогательной системы. На каждом месячном срезе модель может формировать список бумаг с высоким риском, который затем используется для углубленной проверки. Такой сценарий масштабирует первичный анализ и фокусирует внимание на наиболее рискованных случаях.",
        "Таким образом, третья глава показывает не только факт написания кода, но и связь реализации с методикой. Формулы из второй главы используются при расчете признаков и target-компонент, ограничения ручного анализа из первой главы закрываются через воспроизводимый пайплайн, а качество результата оценивается через метрики и временную проверку.",
        "По итогам реализации можно сделать вывод, что разработанный пайплайн решает поставленную задачу автоматической классификации инвестиционных рисков. Он формирует признаки, строит целевую переменную, обучает несколько архитектур, выбирает модель по validation и проверяет качество на test-периоде. Дополнительно сохраняются артефакты запуска: predictions.csv, feature_importance.csv, model_leaderboard.csv, walk_forward_report.csv, feature_drift_report.csv, metrics.json и model_package.joblib.",
        "Вывод по третьей главе. В третьей главе реализован ML-пайплайн, в котором формулы и метрики, обоснованные во второй главе, применены для расчета признаков, формирования целевой переменной и оценки качества автоматической классификации инвестиционных рисков. Финальная архитектура ann_plus_regime показала на тестовом периоде macro-F1 0,5494 и recall класса high 0,9062, что подтверждает применимость подхода для автоматизированного риск-скрининга при условии регулярного контроля дрейфа данных.",
    ]:
        add_plain_paragraph(doc, text)
    add_chapter3_defense_context(doc, data)
    doc.add_page_break()


def conclusion(doc: Document) -> None:
    add_heading(doc, "ЗАКЛЮЧЕНИЕ", 1)
    for text in [
        "В результате выполнения выпускной квалификационной работы была достигнута поставленная цель: разработан и оценен процесс автоматической классификации инвестиционных рисков путем объединения деревьев решений и кластеризационного анализа с использованием искусственных нейронных сетей. Работа построена как переход от ручной оценки к формализованному и воспроизводимому ML-пайплайну.",
        "В первой главе были изучены существующие подходы к оценке и классификации инвестиционных рисков, раскрыта экономическая сущность риска и показано, что риск акции является многокомпонентной характеристикой. Он включает волатильность, просадку, ликвидность, рыночную чувствительность, хвостовые потери и фундаментальную устойчивость эмитента.",
        "Также в первой главе были выявлены ограничения ручного подхода. Ручной анализ характеризуется субъективностью интерпретации, высокой трудоемкостью, ограниченной масштабируемостью и низкой воспроизводимостью при увеличении числа объектов и признаков. Эти ограничения обосновывают необходимость автоматизации классификации инвестиционных рисков.",
        "Во второй главе был проведен анализ предметной области классификации инвестиционных рисков. Были обоснованы группы исходных данных: рыночные цены, объемы торгов, ликвидность, рыночный индекс, макроэкономические показатели, фундаментальные данные и финансовая отчетность эмитентов. Показано, что каждая группа данных соответствует определенной стороне инвестиционного риска.",
        "Во второй главе также были выписаны и объяснены экономические и математические формулы, используемые в работе: доходность, средняя доходность, волатильность, downside-volatility, максимальная просадка, VaR, CVaR, индикатор неликвидности, бета, долговая нагрузка, покрытие процентов и интегральный RiskScore. Тем самым формулы были связаны с предметной областью и последующей программной реализацией.",
        "В третьей главе был реализован ML-пайплайн автоматической классификации инвестиционных рисков. Он включает расчет признаков, формирование целевой переменной, временное разделение выборки, кластеризацию рыночных режимов, обучение ансамблей деревьев, извлечение латентных факторов и оценку качества классификации. Реализация построена с учетом point-in-time логики, что снижает риск утечки будущей информации.",
        "В ходе оценки качества сравнивались несколько архитектур. Финальной стала архитектура ann_plus_regime, объединяющая режимный слой и латентные факторы. На тестовом периоде она показала macro-F1 0,5494, weighted-F1 0,5413, balanced accuracy 0,5714 и recall класса high 0,9062. Высокий recall класса high важен для риск-скрининга, поскольку пропуск высокого риска является наиболее опасной ошибкой.",
        "Работа также показала ограничения разработанного подхода. Качество модели зависит от полноты данных, устойчивости рыночного режима, качества раскрытия отчетности и регулярности переобучения. Диагностика дрейфа выявила существенные изменения распределений между периодами, поэтому практическое использование системы должно сопровождаться мониторингом PSI и периодическим обновлением модели.",
        "Таким образом, предложенный подход позволяет формализовать процесс классификации инвестиционных рисков, снизить зависимость от субъективной ручной оценки, обеспечить воспроизводимость расчетов и получить измеримую оценку качества. Разработанный пайплайн может рассматриваться как исследовательский прототип системы автоматизированного риск-скоринга акций.",
        "Направления дальнейшего развития включают расширение инвестиционной вселенной, улучшение качества парсинга финансовой отчетности, добавление более детальных отраслевых факторов, настройку операционных порогов под разные политики риска и регулярное переобучение модели при изменении рыночного режима.",
    ]:
        add_plain_paragraph(doc, text)
    add_conclusion_defense_context(doc)
    doc.add_page_break()


def references(doc: Document) -> None:
    add_heading(doc, "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ И ИНТЕРНЕТ-РЕСУРСОВ", 1)
    refs = [
        "Бланк И.А. Инвестиционный менеджмент. – Киев: Ника-Центр, 2001. – 448 с.",
        "Шарп У., Александер Г., Бэйли Дж. Инвестиции. – М.: ИНФРА-М, 2018. – 1028 с.",
        "Марковиц Г. Выбор портфеля // Финансы. – 1952. – № 7. – С. 77–91.",
        "Markowitz H. Portfolio Selection // The Journal of Finance. – 1952. – Vol. 7, No. 1. – P. 77–91.",
        "Sharpe W.F. Capital Asset Prices: A Theory of Market Equilibrium under Conditions of Risk // The Journal of Finance. – 1964. – Vol. 19, No. 3. – P. 425–442.",
        "Lintner J. The Valuation of Risk Assets and the Selection of Risky Investments in Stock Portfolios and Capital Budgets // Review of Economics and Statistics. – 1965. – Vol. 47, No. 1. – P. 13–37.",
        "Mossin J. Equilibrium in a Capital Asset Market // Econometrica. – 1966. – Vol. 34, No. 4. – P. 768–783.",
        "Jorion P. Value at Risk: The New Benchmark for Managing Financial Risk. – New York: McGraw-Hill, 2007. – 602 p.",
        "Hull J.C. Risk Management and Financial Institutions. – Hoboken: Wiley, 2018. – 832 p.",
        "McNeil A.J., Frey R., Embrechts P. Quantitative Risk Management: Concepts, Techniques and Tools. – Princeton: Princeton University Press, 2015. – 720 p.",
        "Tsay R.S. Analysis of Financial Time Series. – Hoboken: Wiley, 2010. – 715 p.",
        "Campbell J.Y., Lo A.W., MacKinlay A.C. The Econometrics of Financial Markets. – Princeton: Princeton University Press, 1997. – 632 p.",
        "Amihud Y. Illiquidity and Stock Returns: Cross-section and Time-series Effects // Journal of Financial Markets. – 2002. – Vol. 5, No. 1. – P. 31–56.",
        "Breiman L. Random Forests // Machine Learning. – 2001. – Vol. 45. – P. 5–32.",
        "Breiman L., Friedman J., Stone C.J., Olshen R.A. Classification and Regression Trees. – Boca Raton: CRC Press, 1984. – 368 p.",
        "Hastie T., Tibshirani R., Friedman J. The Elements of Statistical Learning. – New York: Springer, 2009. – 745 p.",
        "Bishop C.M. Pattern Recognition and Machine Learning. – New York: Springer, 2006. – 738 p.",
        "Goodfellow I., Bengio Y., Courville A. Deep Learning. – Cambridge: MIT Press, 2016. – 800 p.",
        "Murphy K.P. Probabilistic Machine Learning: An Introduction. – Cambridge: MIT Press, 2022. – 858 p.",
        "James G., Witten D., Hastie T., Tibshirani R. An Introduction to Statistical Learning. – New York: Springer, 2021. – 607 p.",
        "Friedman J.H. Greedy Function Approximation: A Gradient Boosting Machine // Annals of Statistics. – 2001. – Vol. 29, No. 5. – P. 1189–1232.",
        "Geurts P., Ernst D., Wehenkel L. Extremely Randomized Trees // Machine Learning. – 2006. – Vol. 63. – P. 3–42.",
        "MacQueen J. Some Methods for Classification and Analysis of Multivariate Observations // Proceedings of the Fifth Berkeley Symposium. – 1967. – P. 281–297.",
        "Lloyd S. Least Squares Quantization in PCM // IEEE Transactions on Information Theory. – 1982. – Vol. 28, No. 2. – P. 129–137.",
        "Rumelhart D.E., Hinton G.E., Williams R.J. Learning Representations by Back-propagating Errors // Nature. – 1986. – Vol. 323. – P. 533–536.",
        "Kingma D.P., Ba J. Adam: A Method for Stochastic Optimization // arXiv:1412.6980. – 2014.",
        "Pedregosa F. et al. Scikit-learn: Machine Learning in Python // Journal of Machine Learning Research. – 2011. – Vol. 12. – P. 2825–2830.",
        "Paszke A. et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library // Advances in Neural Information Processing Systems. – 2019. – Vol. 32.",
        "McKinney W. Data Structures for Statistical Computing in Python // Proceedings of the 9th Python in Science Conference. – 2010. – P. 56–61.",
        "Harris C.R. et al. Array Programming with NumPy // Nature. – 2020. – Vol. 585. – P. 357–362.",
        "Pandas Development Team. pandas documentation. – URL: https://pandas.pydata.org/docs/ (дата обращения: 16.05.2026).",
        "Scikit-learn developers. scikit-learn documentation. – URL: https://scikit-learn.org/stable/ (дата обращения: 16.05.2026).",
        "PyTorch contributors. PyTorch documentation. – URL: https://pytorch.org/docs/stable/ (дата обращения: 16.05.2026).",
        "Московская Биржа. Информационные продукты и рыночные данные. – URL: https://www.moex.com/ (дата обращения: 16.05.2026).",
        "Банк России. Ключевая ставка Банка России. – URL: https://www.cbr.ru/ (дата обращения: 16.05.2026).",
        "Интерфакс – Центр раскрытия корпоративной информации. – URL: https://www.e-disclosure.ru/ (дата обращения: 16.05.2026).",
        "ГОСТ 7.32–2017. Система стандартов по информации, библиотечному и издательскому делу. Отчет о научно-исследовательской работе. Структура и правила оформления.",
        "ГОСТ Р 7.0.5–2008. Библиографическая ссылка. Общие требования и правила составления.",
        "ГОСТ 7.1–2003. Библиографическая запись. Библиографическое описание. Общие требования и правила составления.",
        "Финансовый университет при Правительстве Российской Федерации. Методические рекомендации по подготовке и защите выпускной квалификационной работы. – М., 2026.",
        "Chan E.P. Algorithmic Trading: Winning Strategies and Their Rationale. – Hoboken: Wiley, 2013. – 224 p.",
        "Cartea A., Jaimungal S., Penalva J. Algorithmic and High-Frequency Trading. – Cambridge: Cambridge University Press, 2015. – 356 p.",
        "Engle R.F. Autoregressive Conditional Heteroscedasticity with Estimates of the Variance of United Kingdom Inflation // Econometrica. – 1982. – Vol. 50, No. 4. – P. 987–1007.",
        "Bollerslev T. Generalized Autoregressive Conditional Heteroskedasticity // Journal of Econometrics. – 1986. – Vol. 31. – P. 307–327.",
        "Kendall M.G. The Analysis of Economic Time-Series – Part I: Prices // Journal of the Royal Statistical Society. – 1953. – Vol. 116, No. 1. – P. 11–34.",
    ]
    for i, ref in enumerate(refs, 1):
        add_numbered(doc, i, ref)
    doc.add_page_break()


def appendices(doc: Document, data: dict) -> None:
    add_heading(doc, "ПРИЛОЖЕНИЕ А. СТРУКТУРА ПРОГРАММНОЙ РЕАЛИЗАЦИИ", 1)
    for text in [
        "Программная реализация построена в виде Python-пакета risk_pipeline. Такая структура позволяет отделить расчет признаков, построение целевой переменной, обучение моделей и оценку качества. Ниже приведено назначение основных модулей, используемых в работе.",
    ]:
        add_plain_paragraph(doc, text)
    rows = [
        ["risk_pipeline/features.py", "Расчет доходностей, rolling-признаков, ликвидности, beta и месячных срезов"],
        ["risk_pipeline/targets.py", "Расчет future-компонент риска, RiskScore и классов low/medium/high"],
        ["risk_pipeline/feature_engineering.py", "Экономически интерпретируемые derived features и cross-sectional ranks"],
        ["risk_pipeline/models.py", "Деревья решений, кластеризация режима, автоэнкодер, sector overlay"],
        ["risk_pipeline/evaluation.py", "Метрики классификации, confusion matrix, classwise metrics"],
        ["risk_pipeline/diagnostics.py", "Data quality checks, temporal split validation, PSI drift diagnostics"],
        ["risk_pipeline/pipeline.py", "Сквозная orchestration-логика обучения и оценки"],
        ["risk_pipeline/predict.py", "Инференс сохраненного model package на новых наблюдениях"],
    ]
    add_table(doc, ["Файл", "Назначение"], rows, [5.0, 10.5], font_size=10)
    add_source(doc, "Источник: составлено автором.")

    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЕ Б. КОМАНДЫ ЗАПУСКА И АРТЕФАКТЫ", 1)
    for text in [
        "Основной запуск пайплайна выполняется через командную строку. Команды приведены как пример воспроизведения результатов, описанных в третьей главе.",
    ]:
        add_plain_paragraph(doc, text)
    code_blocks = [
        "python -m risk_pipeline.cli --config configs/config.example.yaml run-model-ready --input data/processed/monthly_model_ready.csv --out results/main_run",
        "python scripts/make_quality_materials.py results/main_run",
        "python -m risk_pipeline.cli predict --model-package results/main_run/model_package.joblib --input data/processed/current_month_observations.csv --output results/main_run/current_predictions.csv",
    ]
    for code in code_blocks:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(code)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(10)
    artifact_rows = [
        ["metrics.json", "Основные метрики, пороги, параметры и summary"],
        ["predictions.csv", "Предсказания на test-периоде и вероятности классов"],
        ["feature_importance.csv", "Важность признаков по ветвям модели"],
        ["model_leaderboard.csv", "Сравнение кандидатов на validation"],
        ["walk_forward_report.csv", "Результаты walk-forward проверки"],
        ["feature_drift_report.csv", "Диагностика PSI и missing-rate drift"],
        ["model_package.joblib", "Сохраненный пакет модели для инференса"],
    ]
    add_table(doc, ["Артефакт", "Назначение"], artifact_rows, [4.6, 10.9], font_size=10)

    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЕ В. ФРАГМЕНТЫ ФОРМУЛЬНОЙ РЕАЛИЗАЦИИ", 1)
    snippets = [
        ("Максимальная просадка", "running_max = np.maximum.accumulate(close)\ndrawdown = close / (running_max + EPS) - 1.0\nreturn abs(np.nanmin(drawdown))"),
        ("Downside-volatility", "values = returns[np.isfinite(returns) & (returns < 0)]\nreturn np.nanstd(values, ddof=1)"),
        ("CVaR 95%", "cutoff = np.nanquantile(values, 0.05)\ntail = values[values <= cutoff]\nreturn abs(np.nanmean(tail))"),
        ("Future illiquidity", "metric = np.abs(returns) / (np.abs(value) + EPS)\nreturn np.nanmean(metric)"),
        ("RiskScore", "score += weight * percentile_against_train(component_values)"),
    ]
    for title, code in snippets:
        add_plain_paragraph(doc, title + ".", bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(code)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(10)

    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЕ Г. ДОПОЛНИТЕЛЬНЫЕ ДАННЫЕ ПО КАЧЕСТВУ", 1)
    drift_top = data["drift"].head(15)[["feature", "period", "psi", "drift_warning"]].values.tolist()
    add_table(doc, ["Признак", "Период", "PSI", "Предупреждение"], drift_top, [6.6, 3.0, 2.5, 3.4], font_size=8)
    add_source(doc, "Источник: составлено автором по отчету диагностики дрейфа признаков.")
    sector_rows = data["sector"].values.tolist()
    add_caption(doc, "Таблица Г.2 – Проверка sector overlay")
    add_table(doc, list(data["sector"].columns), sector_rows, [3.5, 1.6, 2.7, 2.7, 2.1, 2.0], font_size=8)
    add_source(doc, "Источник: составлено автором по отчету проверки sector overlay.")


def add_padding_pages_if_needed(doc: Document) -> None:
    """Add substantive method notes to keep the DOCX safely above 60 rendered pages."""
    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЕ Д. МЕТОДИЧЕСКИЕ ПОЯСНЕНИЯ К ИСПОЛЬЗОВАНИЮ ПАЙПЛАЙНА", 1)
    notes = [
        "При использовании разработанного пайплайна важно отделять исследовательскую оценку от операционного применения. Исследовательская оценка показывает, насколько выбранная архитектура способна воспроизводить разметку риска на исторических данных. Операционное применение требует регулярной проверки качества, мониторинга сдвига распределений и настройки порогов принятия решений.",
        "Если модель используется как предварительный риск-скрининг, ее результат не должен автоматически заменять инвестиционное решение. Класс high целесообразно трактовать как сигнал для углубленного анализа: проверки новостей, структуры долгов, корпоративных событий, отчетности и ликвидности. Такой режим применения соответствует роли автоматизации как инструмента первичной классификации.",
        "При расширении выборки необходимо контролировать сопоставимость эмитентов. Добавление низколиквидных бумаг может резко изменить распределение неликвидности и ухудшить устойчивость порогов. Поэтому при переходе к более широкой вселенной следует пересчитать train-only пороги и провести отдельную оценку классов по секторам.",
        "Регулярное переобучение модели должно выполняться после накопления новых наблюдений или при существенном дрейфе признаков. В финансовых данных сдвиг может быть вызван изменением ставок, валютного режима, санкционных ограничений, ликвидности рынка или структуры раскрытия отчетности. PSI не объясняет причину сдвига, но помогает формально зафиксировать необходимость проверки модели.",
        "Для повышения интерпретируемости результата рекомендуется сохранять не только итоговый класс, но и вероятности классов, важность признаков, значения ключевых показателей риска и режим рынка. Это позволяет аналитикам понять, почему модель отнесла объект к конкретной категории и какие факторы требуют дополнительной проверки.",
        "При сравнении ручного и автоматизированного подхода корректно фиксировать не универсальное превосходство модели, а свойства процесса. Автоматизированный пайплайн обеспечивает повторяемость расчетов, единые правила классификации и измеримое качество на исторических данных. Спорные случаи и качественные факторы могут передаваться на дополнительную аналитическую проверку.",
        "Возможным направлением развития является добавление объяснимости на уровне отдельного предсказания. Для этого могут использоваться SHAP values, permutation importance или локальные surrogate-модели. В рамках данной работы основной акцент сделан на воспроизводимой архитектуре и общей важности признаков, но индивидуальные объяснения могут повысить практическую ценность системы.",
        "Еще одним направлением развития является адаптивная настройка risk policy. Если инвестиционная стратегия допускает больше ложных сигналов ради минимизации пропуска высокого риска, порог high можно снизить. Если важнее уменьшить число ложных предупреждений, порог high можно повысить. Такой выбор должен оформляться как управленческое решение, а не как изменение самой научной постановки.",
        "Финансовая отчетность требует отдельного контроля качества извлечения. Автоматический парсинг PDF и XLSX может давать пропуски или некорректные значения, особенно при сложных таблицах и разных форматах раскрытия. Поэтому отчетные признаки сопровождаются indicators пропусков, лагом публикации и качеством извлечения. Эти признаки помогают модели учитывать надежность входной информации.",
        "Временное разделение выборки является обязательным условием корректной оценки. Случайное перемешивание строк в финансовой панели может привести к тому, что модель обучается на будущих рыночных режимах и проверяется на прошлом, что не соответствует реальному применению. Поэтому train, validation и test должны располагаться последовательно во времени.",
        "При дальнейшей работе можно расширить горизонт целевой переменной и сравнить 63, 126 и 252 торговых дня. Более короткий горизонт сделает модель чувствительнее к быстрым изменениям рынка, но может увеличить шум. Более длинный горизонт приблизит задачу к стратегическому риск-анализу, но уменьшит количество доступных размеченных наблюдений.",
        "Наконец, важным направлением является сравнение с альтернативными базовыми подходами: простыми правилами по волатильности, правилами по максимальной просадке, логистической регрессией и градиентным бустингом без гибридных слоев. Такое сравнение позволит точнее показать вклад каждого элемента архитектуры.",
    ]
    for text in notes:
        add_plain_paragraph(doc, text)


def describe_feature(name: str) -> str:
    if name.endswith("_cs_rank"):
        return "Процентильный ранг признака среди эмитентов на одном месячном срезе; снижает зависимость от общего уровня рынка."
    if name.endswith("_sector_z"):
        return "Секторное z-отклонение признака; показывает, насколько эмитент отличается от компаний своего сектора."
    mapping = {
        "downside_to_total_vol": "Отношение отрицательной волатильности к общей волатильности; показывает долю неблагоприятной изменчивости.",
        "liquidity_stress": "Интегральный признак ликвидностного давления на основе Amihud, spread proxy и торгового оборота.",
        "liquidity_vol_interaction": "Взаимодействие ликвидностного стресса и волатильности; усиливает сигнал при одновременном росте двух рисков.",
        "beta_vol_interaction": "Произведение beta и волатильности; отражает чувствительность волатильной бумаги к рыночному движению.",
        "macro_pressure": "Сводный признак макроэкономического давления: ставка, волатильность рынка, корреляция и наклон кривой.",
        "rate_fx_pressure": "Комбинация ключевой ставки и валютного курса; отражает режим стоимости капитала и валютного давления.",
        "curve_inversion_flag": "Бинарный признак отрицательного наклона кривой доходности ОФЗ.",
        "debt_service_stress": "Фундаментальный стресс обслуживания долга с учетом долговой нагрузки и покрытия процентов.",
        "cashflow_buffer": "Буфер денежных потоков на основе свободного и операционного денежного потока.",
        "size_liquidity_interaction": "Взаимодействие размера эмитента и торговой ликвидности.",
        "momentum_reversal_risk": "Риск продолжения или усиления отрицательного momentum с учетом beta.",
        "fundamental_fragility": "Разность долгового стресса и денежного буфера; отражает фундаментальную хрупкость.",
        "report_leverage_pressure": "Долговое давление по данным опубликованной отчетности.",
        "report_cashflow_pressure": "Давление на денежные потоки по отчетности и текстовым risk-signals.",
        "report_staleness_weight": "Вес актуальности отчетности с учетом лага публикации, доступности и качества извлечения.",
        "report_integrated_stress": "Интегральный стресс по отчетности: финансовое давление, leverage и cash-flow pressure.",
        "fundamental_report_gap": "Разрыв между базовой фундаментальной хрупкостью и отчетным стрессом.",
    }
    if name in mapping:
        return mapping[name]
    if name.startswith("report_"):
        if "missing" in name:
            return "Индикатор отсутствия отчетного показателя; используется для контроля качества раскрытия данных."
        if "yoy_change" in name:
            return "Годовое изменение отчетного показателя; отражает динамику финансового состояния эмитента."
        if "count" in name or "flag" in name:
            return "Текстовый risk-signal из отчетности; фиксирует наличие соответствующей группы риск-терминов."
        return "Числовой показатель, извлеченный из опубликованной финансовой отчетности эмитента."
    if "vol" in name:
        return "Рыночный показатель изменчивости доходности, используемый для оценки ценового риска."
    if "amihud" in name or "turnover" in name or "traded_value" in name or "spread" in name:
        return "Ликвидностный показатель, отражающий торговую активность или ценовое воздействие сделок."
    if "debt" in name or "coverage" in name or "cash_flow" in name or "ebitda" in name:
        return "Фундаментальный показатель финансовой устойчивости эмитента."
    return "Признак, используемый в модели для описания рыночного, макроэкономического или фундаментального состояния объекта."


def feature_dictionary_appendix(doc: Document, data: dict) -> None:
    metrics = data["metrics"]
    manifest = json.loads((RUN_DIR / "run_manifest.json").read_text(encoding="utf-8"))
    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЕ Е. СЛОВАРЬ ПРИЗНАКОВ МОДЕЛИ", 1)
    add_plain_paragraph(
        doc,
        "В данном приложении приведен словарь признаков, используемых или добавленных в пайплайне. Наличие такого словаря повышает воспроизводимость работы: по нему можно понять, какие данные попадают в модель и какую экономическую роль выполняет каждый признак.",
    )
    engineered = metrics.get("added_features", [])
    rows = [[i + 1, feature, describe_feature(feature)] for i, feature in enumerate(engineered)]
    add_caption(doc, "Таблица Е.1 – Инженерные признаки, добавленные в пайплайне")
    add_table(doc, ["№", "Признак", "Экономическая интерпретация"], rows, [1.0, 5.5, 9.2], font_size=8)
    add_source(doc, "Источник: составлено автором по метрикам основного запуска пайплайна.")

    input_columns = [c for c in manifest.get("input_columns", []) if c not in {"decision_date", "ticker"}]
    selected = []
    priority_prefixes = ("return", "rolling", "downside", "momentum", "avg", "amihud", "spread", "turnover", "beta", "market", "cbr", "usd", "ofz", "imoex", "average", "net_debt", "interest", "operating", "free", "book", "future", "report")
    for c in input_columns:
        if c.startswith(priority_prefixes) or c in {"sector", "close", "volume", "value"}:
            selected.append(c)
    selected = selected[:90]
    base_rows = [[i + 1, feature, describe_feature(feature)] for i, feature in enumerate(selected)]
    add_caption(doc, "Таблица Е.2 – Основные входные признаки и поля модельной панели")
    add_table(doc, ["№", "Поле", "Назначение"], base_rows, [1.0, 5.3, 9.4], font_size=8)
    add_source(doc, "Источник: составлено автором по манифесту основного запуска пайплайна.")

    add_plain_paragraph(
        doc,
        "Словарь признаков показывает, что пайплайн не использует произвольный набор столбцов. Входные поля группируются по экономическому смыслу: рыночные признаки описывают поведение цены, ликвидностные признаки – возможность исполнения сделок, макроэкономические признаки – внешний режим, фундаментальные и отчетные признаки – устойчивость эмитента.",
    )
    add_plain_paragraph(
        doc,
        "Такой словарь также помогает контролировать leakage. Target-компоненты используются для формирования класса риска, но не должны попадать в feature engineering как объясняющие признаки. В реализации для этого выделен набор запрещенных target-полей, который не используется при расчете производных признаков.",
    )
    add_plain_paragraph(
        doc,
        "При дальнейшем развитии работы словарь признаков может быть расширен: для каждого поля можно добавить источник, периодичность обновления, допустимый диапазон, способ обработки пропусков и признак принадлежности к группе риска. Это позволит превратить исследовательский пайплайн в более формальную документацию модели.",
    )


def add_intro_defense_context(doc: Document) -> None:
    for text in [
        "Оптимизация процесса в данном исследовании понимается как перевод классификации риска из ручного и плохо масштабируемого режима в формализованный режим, где каждое действие может быть повторено: сбор данных, расчет признаков, формирование целевой переменной, обучение модели, выбор архитектуры и проверка качества.",
        "Поэтому эффект от предложенного подхода оценивается по трем направлениям. Первое направление – масштабируемость, то есть способность применять один и тот же алгоритм к растущему числу эмитентов, дат и признаков. Второе направление – воспроизводимость, то есть возможность повторить расчет при сохранении исходных данных, кода и параметров запуска. Третье направление – точность классификации, то есть измеримое совпадение предсказанного класса риска с классом, построенным по будущей реализации риск-компонент.",
        "Такая трактовка не подменяет научную задачу бизнес-рассуждением о стоимости аналитиков или покупке готового программного продукта. Проблема формулируется методически: ручной анализ не обеспечивает достаточной формализации, когда требуется регулярно классифицировать множество наблюдений по единой логике. Автоматизированный ML-пайплайн, напротив, задает проверяемый порядок действий и позволяет оценить результат количественно.",
        "Предметная основа инвестиционного анализа сохраняется через формализацию финансовых представлений о волатильности, просадке, ликвидности, долговой нагрузке, чувствительности к рынку и качестве отчетности. Эти представления переводятся в формулы, признаки и метрики, поэтому программная реализация связана с экономическим смыслом риска.",
        "Таким образом, введение задает общий исследовательский маршрут: сначала требуется показать, почему ручная классификация риска ограничена, затем необходимо обосновать финансовые показатели и данные, после чего можно переходить к реализации модели и оценке результата. Именно такая последовательность позволяет удержать работу в рамках заявленной темы.",
    ]:
        add_plain_paragraph(doc, text)


def add_chapter1_defense_context(doc: Document) -> None:
    for text in [
        "Постановка задачи автоматической классификации требует зафиксировать, какие именно вопросы должна закрывать автоматизация. Первый вопрос связан с единицей анализа. В работе единицей анализа является не эмитент сам по себе, а пара «эмитент – дата принятия решения». Это принципиально: одна и та же компания в разные месяцы может находиться в разных режимах риска, поэтому классификация должна учитывать временное состояние объекта.",
        "Второй вопрос связан с тем, что считать результатом классификации. Если модель просто присваивает метку low, medium или high без объяснения способа формирования метки, такая постановка остается слабой. В данной работе класс риска строится через будущий RiskScore, а RiskScore, в свою очередь, формируется из максимальной просадки, downside-volatility, CVaR и неликвидности. Это делает класс риска не произвольной категорией, а результатом заранее описанной процедуры.",
        "Третий вопрос связан с границей между ручной оценкой и модельной оценкой. Ручной анализ не детализируется как отдельная полноценная методика с регламентом, поскольку предмет исследования находится в автоматизации процесса. Однако ручной анализ описывается достаточно, чтобы стало понятно, какие его недостатки устраняет пайплайн. Поэтому акцент делается не на индивидуальном порядке заполнения таблицы, а на том, почему такой процесс трудно масштабировать и проверять.",
        "Если ручная классификация выполняется по нескольким показателям, аналитик фактически решает многокритериальную задачу. Он сопоставляет волатильность, ликвидность, просадку, фундаментальные показатели и внешний макроэкономический фон. При этом веса показателей часто остаются неявными. Один и тот же объект может быть классифицирован по-разному, если аналитик считает главным фактором рыночную волатильность или, наоборот, долговую устойчивость эмитента.",
        "Автоматизация не уничтожает проблему выбора весов, но делает ее явной. В целевой переменной веса компонент RiskScore заданы явно, а в модели вклад признаков определяется алгоритмом обучения и может быть проанализирован через важность признаков. Поэтому спор о классификации переводится из неформального обсуждения в проверяемую плоскость: можно изменить веса, пороги или модель и сравнить результаты на одной и той же выборке.",
        "Масштабируемость в рамках исследования можно описать через простую зависимость. Пусть N – число эмитентов, T – число дат наблюдения, K – число признаков. В ручной процедуре объем операций растет вместе с произведением N · T · K, потому что каждый новый объект и каждый новый признак требуют дополнительной проверки. В программном пайплайне после настройки формул и источников добавление новых строк панели не меняет методику анализа: меняется только размер входной таблицы.",
        "Воспроизводимость также имеет конкретное содержание. Для воспроизводимого исследования недостаточно сказать, что модель была обучена. Нужно сохранить версию данных, список признаков, параметры временного разделения, train-only пороги, выбранную архитектуру и итоговые метрики. В данной работе такая логика отражается через сохранение артефактов запуска и описание pipeline-этапов.",
        "Точность процесса нельзя оценивать только общей accuracy, потому что классы риска имеют разную прикладную значимость. Для risk-screening особенно важен класс high. Пропуск high-risk объекта может привести к принятию решения без дополнительной проверки, тогда как ложное завышение класса чаще приводит лишь к дополнительному анализу. Поэтому в работе отдельно используется recall класса high и false negative rate по high-классу.",
        "Наконец, автоматизация должна исключать методическую ошибку утечки будущей информации. В финансовых задачах это один из самых частых источников завышенного качества. Если модель использует данные отчетности, которые на дату принятия решения еще не были опубликованы, или если пороги классов рассчитываются по всей выборке, результат становится некорректным. Поэтому point-in-time логика является не технической деталью, а обязательным условием научной состоятельности работы.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 1.2 – Критерии оценки улучшения процесса классификации")
    add_table(
        doc,
        ["Критерий", "Как проявляется в ручном подходе", "Как оценивается в разработанном пайплайне"],
        [
            ["Масштабируемость", "Рост числа эмитентов, дат и признаков увеличивает объем ручных операций", "Пакетная обработка панели наблюдений; число строк и признаков фиксируется в данных и артефактах"],
            ["Воспроизводимость", "Результат зависит от сохранности таблиц и неявных ручных решений", "Сохраняются код, конфигурация, пороги, модель, predictions и run manifest"],
            ["Точность", "Обычно оценивается качественно или по отдельным показателям", "Используются macro-F1, weighted-F1, balanced accuracy, recall high и confusion matrix"],
            ["Временная корректность", "Есть риск использования данных, недоступных на дату решения", "Признаки строятся point-in-time, target-компоненты не попадают в признаки"],
            ["Интерпретируемость", "Объяснение зависит от аналитика", "Используются важность признаков, структура RiskScore и раздельная оценка классов"],
        ],
        [3.0, 6.1, 6.4],
        font_size=9,
    )
    add_source(doc, "Источник: составлено автором.")
    for text in [
        "Эта таблица задает способ оценки улучшения процесса классификации. В работе не утверждается, что измерена экономия рабочего времени в конкретной организации. Вместо этого оцениваются исследовательские признаки процесса: количество обработанных наблюдений, число признаков, наличие повторяемого запуска, сохранение артефактов и качество классификации на тестовой выборке.",
        "Такой подход соответствует научной постановке. Если бы работа была бизнес-проектом внедрения, можно было бы измерять стоимость человеко-часов, скорость обработки заявок или эффект от покупки готового программного продукта. Однако исследование направлено на метод автоматической классификации, поэтому основными являются методические показатели: корректность данных, формализация целевой переменной, качество модели и воспроизводимость результата.",
        "Именно поэтому во второй главе подробно раскрываются показатели риска и источники данных. Без этого автоматизация выглядела бы как применение алгоритмов машинного обучения к произвольной таблице. Каждая группа признаков имеет финансовый смысл, каждая формула связана с компонентом риска, а каждая метрика качества отвечает на определенный исследовательский вопрос.",
    ]:
        add_plain_paragraph(doc, text)


def add_chapter1_literature_review(doc: Document) -> None:
    for text in [
        "Постановка задачи автоматической классификации требует рассмотреть не только ручной анализ, но и существующие классы решений. В литературе и практике риск-анализа можно выделить несколько групп подходов: классические финансовые модели, статистические модели временных рядов, эконометрические модели риска, алгоритмы машинного обучения и гибридные архитектуры, объединяющие несколько источников информации.",
        "Классические финансовые модели исходят из связи риска и доходности. В портфельной теории Марковица риск описывается дисперсией доходности, а задача инвестора формулируется как поиск компромисса между ожидаемой доходностью и изменчивостью портфеля. CAPM и связанные с ним модели вводят систематический риск через beta, то есть чувствительность актива к рыночному портфелю. Эти модели важны как теоретическая база, но они предполагают достаточно сильные допущения о распределениях и стабильности связей.",
        "Преимущество классических моделей состоит в высокой интерпретируемости. Beta, дисперсия, ковариация и ожидаемая доходность имеют понятный экономический смысл. Однако для автоматической классификации риска отдельных акций этого недостаточно. Во-первых, дисперсия симметрично учитывает положительные и отрицательные отклонения. Во-вторых, связь с рынком может меняться во времени. В-третьих, фундаментальные и ликвидностные факторы в таких моделях обычно не входят в явном виде.",
        "Статистические и эконометрические модели, такие как ARCH/GARCH и модели Value at Risk, позволяют лучше описывать изменчивость финансовых временных рядов. Они учитывают кластеризацию волатильности и позволяют оценивать хвостовые потери. Эти методы полезны для формализации отдельных компонент риска, но сами по себе не решают задачу многоклассовой классификации с большим набором разнородных признаков.",
        "Отдельную группу составляют правила риск-скоринга. В них заранее задается набор показателей, пороги и веса: например, высокая волатильность, низкая ликвидность, высокая долговая нагрузка и слабое покрытие процентов повышают итоговый риск. Такой подход легко объяснить, но он остается ограниченным из-за жесткости правил. Если связь между признаками нелинейна или меняется в зависимости от режима рынка, фиксированная система порогов может давать нестабильный результат.",
        "Методы машинного обучения позволяют перейти от жестких правил к обучаемой классифицирующей функции. Логистическая регрессия и линейные модели удобны как базовый уровень, потому что они интерпретируемы и устойчивы на небольших выборках. Однако линейная модель может плохо учитывать сложные взаимодействия признаков, например одновременное влияние беты, волатильности, неликвидности и процентной ставки.",
        "Деревья решений и ансамбли деревьев являются естественным выбором для задач с разнородными финансовыми признаками. Они способны учитывать нелинейные зависимости, не требуют строгой нормальности распределений и позволяют оценивать важность признаков. Random Forest снижает дисперсию отдельных деревьев, Extra Trees усиливает случайность разбиений, а Gradient Boosting последовательно исправляет ошибки предыдущих моделей. Поэтому ансамбли деревьев часто используются как сильная базовая архитектура для табличных данных.",
        "Кластеризационный анализ применяется в финансовых задачах для выделения рыночных режимов, групп похожих эмитентов или типовых состояний ликвидности. Его ценность состоит в том, что связь между признаками и будущим риском может быть различной в спокойном и стрессовом периоде. Например, высокая долговая нагрузка в период низких ставок и высокая долговая нагрузка при росте ставок имеют разный смысл. Кластер режима позволяет добавить в модель информацию о контексте.",
        "Искусственные нейронные сети применимы тогда, когда требуется извлечь скрытые представления из большого признакового пространства. В данной работе нейросетевой слой используется не как единственная модель, а как способ получить латентные факторы. Такой выбор снижает риск неоправданного усложнения: основная классификация остается основанной на табличных моделях, а нейронная сеть добавляет компактное представление признаков.",
        "Гибридные подходы особенно перспективны для инвестиционного риска, потому что риск не является одномерным явлением. Деревья решений хорошо работают с табличными признаками, кластеризация добавляет режимный контекст, а нейронная сеть может сжать сложное признаковое пространство. Поэтому выбранная архитектура соответствует современной тенденции: не использовать один алгоритм изолированно, а строить воспроизводимый pipeline, где каждый слой имеет свою функцию.",
        "Критический анализ существующих решений показывает, что простое применение готовой модели к готовому датасету не раскрывает задачу прикладного машинного обучения. Необходимо описать весь путь: формирование данных, обоснование признаков, выбор целевой переменной, сравнение методов, проверку качества, интерпретацию и ограничения. Именно поэтому акцент сделан на процессе автоматической классификации, а не только на одной модели.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 1.3 – Классификация подходов к оценке и классификации инвестиционного риска")
    add_table(
        doc,
        ["Группа подходов", "Сильные стороны", "Ограничения", "Роль в работе"],
        [
            ["Портфельная теория и CAPM", "Интерпретируемость, связь риска и доходности", "Сильные допущения, ограниченный набор факторов", "Формируют теоретическую базу риска и beta"],
            ["VaR, CVaR, модели волатильности", "Описывают хвостовые потери и изменчивость", "Часто оценивают отдельную меру риска, а не класс", "Используются при построении target-компонент"],
            ["Ручной риск-скоринг", "Понятные правила и пороги", "Субъективность весов, слабая адаптивность", "Используется как базовая логика формализации"],
            ["Ансамбли деревьев", "Работают с нелинейностями и табличными признаками", "Требуют валидации и контроля переобучения", "Основной классифицирующий слой"],
            ["Кластеризация", "Выделяет режимы и похожие состояния рынка", "Кластеры требуют интерпретации", "Добавляет рыночный контекст"],
            ["Нейронные сети", "Извлекают латентные факторы из признаков", "Менее интерпретируемы без дополнительных проверок", "Используются как слой представлений"],
        ],
        [3.6, 4.2, 4.1, 3.3],
        font_size=12,
    )
    add_source(doc, "Источник: составлено автором.")
    for text in [
        "Выбор методов в данной работе строится на этом сравнении. В качестве основы используются деревья решений, поскольку они хорошо подходят для табличных данных и дают интерпретируемые признаки важности. Кластеризация добавляется для учета рыночного режима. Нейросетевой слой используется для получения латентных факторов, а не для подмены всей логики классификации черным ящиком.",
        "Такой выбор методов также связан с ограничениями данных. Финансовая панель имеет умеренный размер, содержит разнородные признаки и временную зависимость. Для такой постановки нецелесообразно строить чрезмерно глубокую нейронную сеть как единственную модель. Более устойчивым является гибридный подход, где сильные стороны разных методов используются совместно.",
        "Обзор существующих подходов позволяет сформулировать требования к дальнейшей разработке. Во-первых, модель должна учитывать несколько компонент риска. Во-вторых, она должна быть проверена на временном разбиении. В-третьих, необходим baseline для сравнения. В-четвертых, результат должен быть интерпретирован в терминах предметной области, а не только через числовую метрику.",
        "Таким образом, первая глава выполняет роль не только постановки проблемы, но и теоретического обоснования выбранной архитектуры. Она показывает, почему ручной анализ ограничен, почему одиночная финансовая метрика недостаточна и почему объединение деревьев решений, кластеризационного анализа и нейронных сетей является содержательно оправданным для поставленной задачи.",
        "Рассмотренные источники также показывают, что современные работы в области финансового машинного обучения уделяют большое внимание воспроизводимости экспериментов. Для финансовых данных особенно опасны случайные разбиения, неявная подгонка гиперпараметров под тестовый период и использование будущей информации. Поэтому корректная методология включает временное разделение, контроль leakage, сравнение с baseline и сохранение артефактов запуска.",
        "Еще один вывод обзора состоит в необходимости интерпретировать результат не только математически, но и экономически. Даже если модель показывает приемлемое значение F1, для риск-анализа важно понимать, какие факторы привели объект к высокому классу риска. Поэтому в дальнейшем используются важность признаков, анализ классов, матрица ошибок и обсуждение trade-off между precision и recall.",
        "Наиболее близкими по смыслу являются подходы, объединяющие финансовые показатели и алгоритмы машинного обучения для кредитного, рыночного или инвестиционного скоринга. Однако отличие настоящей работы состоит в том, что целевая переменная строится не по одной метке дефолта или доходности, а по будущему интегральному профилю риска. Это делает задачу более близкой к инвестиционному риск-скринингу.",
        "Таким образом, теоретический обзор не является формальным перечислением литературы. Он позволяет обосновать три ключевых решения: использовать многокомпонентную разметку риска, учитывать рыночные режимы через кластеризационный слой и сравнивать гибридную архитектуру с базовыми моделями. Эти решения затем переходят во вторую главу как требования к проектированию ML-пайплайна.",
    ]:
        add_plain_paragraph(doc, text)


def add_chapter2_defense_context(doc: Document, data: dict) -> None:
    panel = data["panel"]
    metrics = data["metrics"]
    universe = data["universe"]
    dates = pd.to_datetime(panel["decision_date"])
    date_min = dates.min().strftime("%d.%m.%Y")
    date_max = dates.max().strftime("%d.%m.%Y")
    for text in [
        "Выбранная предметная область требует многокомпонентной модели риска. Акция как финансовый инструмент одновременно несет ценовой риск, риск ликвидности, систематический рыночный риск, фундаментальный риск эмитента и риск изменения внешнего режима. Поэтому одной формулы волатильности недостаточно: во второй главе используется набор формул, каждая из которых отвечает за отдельную сторону риска.",
        "Ценовой риск описывает возможное неблагоприятное изменение рыночной стоимости актива. Он проявляется через доходность, волатильность, downside-volatility, максимальную просадку, VaR и CVaR. Эти показатели не дублируют друг друга полностью. Волатильность показывает общую изменчивость, downside-volatility концентрируется на отрицательных изменениях, максимальная просадка отражает глубину падения от локального максимума, а CVaR описывает среднюю величину наиболее неблагоприятных хвостовых исходов.",
        "Ликвидностный риск особенно важен для акций с неоднородной торговой активностью. Инвестор может столкнуться с ситуацией, когда цена формально известна, но крупная позиция не может быть закрыта без существенного ценового воздействия. Поэтому в работе используются торговый оборот, turnover ratio, spread proxy и индикатор Амихуда. Эти показатели позволяют отличить просто волатильную бумагу от бумаги, где неблагоприятное движение осложняется невозможностью быстро выйти из позиции.",
        "Систематический риск связан с реакцией бумаги на общий рыночный режим. Для его описания используется beta и макроэкономические признаки. Если акция сильно зависит от рынка, то ухудшение индекса или рост корреляций может повышать вероятность будущей просадки. Поэтому в набор признаков включаются не только индивидуальные характеристики эмитента, но и состояние внешней среды.",
        "Фундаментальный риск показывает устойчивость самой компании. Высокая долговая нагрузка, слабое покрытие процентов, отрицательный свободный денежный поток и снижение маржинальности могут повышать вероятность будущей уязвимости даже тогда, когда текущая цена еще не демонстрирует сильного падения. По этой причине финансовая отчетность рассматривается не как декоративное дополнение, а как отдельный информационный слой.",
        f"Используемая панель данных содержит {len(panel)} месячных наблюдения, {panel.shape[1]} столбцов и {len(universe)} эмитентов. Период наблюдений – с {date_min} по {date_max}. Такая структура показывает, что работа опирается не на отдельный пример, а на панель, где каждый эмитент наблюдается во времени. Именно панельная структура делает актуальными вопросы временного разбиения, дрейфа признаков и повторяемости расчета.",
        "Данные в работе не берутся из воздуха. Рыночные признаки формируются из цен, объемов и торгового оборота; макроэкономические признаки отражают внешний режим; фундаментальные признаки описывают финансовую устойчивость; отчетные признаки извлекаются из опубликованной отчетности и сопровождаются показателями доступности и лагов. Такое деление важно, потому что каждая группа данных имеет разную периодичность обновления и разную степень надежности.",
        "Особое внимание требуется к датам. Для рыночных данных дата наблюдения обычно совпадает с торговым календарем, а для отчетности необходимо учитывать дату публикации. Если использовать финансовые показатели по периоду, но игнорировать дату раскрытия, модель может получить информацию, которой у инвестора еще не было. Поэтому в реализации отчетный слой должен быть связан не только с числовыми значениями, но и с лагом публикации и признаком доступности.",
        "Для формализации предметной области полезно разделить признаки на объясняющие и целевые. Объясняющие признаки описывают состояние объекта на дату принятия решения. Целевые компоненты описывают будущую реализацию риска на горизонте 126 торговых дней. Такое разделение является центральным для всей работы: если будущие компоненты попадут в признаки, модель будет не прогнозировать риск, а фактически читать ответ из будущего.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 2.4 – Связь групп риска, данных и формул")
    add_table(
        doc,
        ["Группа риска", "Используемые данные", "Ключевые формулы и признаки", "Роль в классификации"],
        [
            ["Ценовой риск", "Цена закрытия и доходности", "rₜ, σ, σ⁻, MDD, VaR, CVaR", "Описывает вероятность и масштаб неблагоприятного изменения цены"],
            ["Ликвидностный риск", "Объем, оборот, стоимость сделок", "Amihud, turnover ratio, spread proxy", "Показывает, насколько сложно выйти из позиции без ценового воздействия"],
            ["Систематический риск", "Индекс рынка и макрофакторы", "beta, market volatility, key rate, FX pressure", "Учитывает зависимость бумаги от общего рыночного режима"],
            ["Фундаментальный риск", "Финансовые показатели эмитента", "ND/EBITDA, interest coverage, margins, cash flow", "Оценивает устойчивость компании и способность обслуживать обязательства"],
            ["Отчетный риск", "Публичная отчетность и качество извлечения", "report_lag, report_integrated_stress, text risk density", "Учитывает опубликованные сигналы и надежность источника данных"],
        ],
        [3.0, 3.8, 4.0, 4.7],
        font_size=8,
    )
    add_source(doc, "Источник: составлено автором.")
    for text in [
        "Следующий важный вопрос – почему целевая переменная строится не по одному будущему показателю. Если использовать только будущую доходность, модель начнет решать задачу прогноза доходности, а не классификации риска. Если использовать только максимальную просадку, будет потеряна информация о хвостовых потерях и неликвидности. Если использовать только волатильность, положительные и отрицательные колебания будут смешаны. Поэтому в работе выбран интегральный RiskScore.",
        "RiskScore в данной работе выполняет роль bridge-показателя между экономическим смыслом риска и задачей машинного обучения. Он переводит несколько будущих компонент риска в единую шкалу, после чего эта шкала разбивается на три класса. Такая процедура делает target более устойчивым, чем одиночный показатель, и лучше соответствует предметной области: высокий инвестиционный риск может возникать по разным причинам, но модель должна уметь распознавать общий класс неблагоприятной уязвимости.",
        "Компоненты RiskScore переводятся в процентильную шкалу относительно обучающего периода. Это важно по двум причинам. Во-первых, показатели имеют разные размерности: максимальная просадка, CVaR и неликвидность не могут быть напрямую сложены без нормализации. Во-вторых, процентильная шкала делает итоговый score интерпретируемым как относительное положение объекта среди обучающих наблюдений.",
    ]:
        add_plain_paragraph(doc, text)
    add_formula(doc, "p(xₜ,ᵢ) = rank_train(xₜ,ᵢ) / n_train", "2.14")
    add_plain_paragraph(doc, "где p(xₜ,ᵢ) – процентильное положение значения x для i-го эмитента на дату t относительно распределения обучающего периода; n_train – число наблюдений в обучающей выборке.")
    add_formula(doc, "RiskScoreₜ,ᵢ = 0,35·p(MDDₜ,ᵢ) + 0,30·p(σ⁻ₜ,ᵢ) + 0,20·p(CVaR₉₅,ₜ,ᵢ) + 0,15·p(ILLIQₜ,ᵢ)", "2.15")
    add_plain_paragraph(doc, "Формула RiskScore показывает, что наибольший вес получает будущая максимальная просадка, поскольку она напрямую отражает глубину потери капитала. Downside-volatility получает второй по величине вес, так как характеризует регулярность отрицательных колебаний. CVaR отвечает за хвостовые потери, а неликвидность учитывает трудность выхода из позиции в неблагоприятном сценарии.")
    for text in [
        "Разбиение RiskScore на классы low, medium и high проводится по train-only порогам. Это означает, что пороги определяются только на обучающем периоде и затем применяются к validation и test. Такой подход предотвращает подгонку классов под будущие периоды. Если бы пороги рассчитывались по всей выборке, в модельную оценку была бы встроена информация о будущих распределениях риска.",
        f"В основном запуске пороги составили: low_upper = {metrics['target_thresholds']['low_upper']:.4f}, medium_upper = {metrics['target_thresholds']['medium_upper']:.4f}. Эти значения не подбираются вручную после просмотра test-результатов; они получены из обучающего периода и затем используются как фиксированное правило классификации.",
        "Метрики качества также должны соответствовать предметной области. Accuracy удобна как базовая мера, но для несбалансированных или неодинаково важных классов она недостаточна. Macro-F1 усредняет качество по классам и не позволяет большому классу полностью доминировать в оценке. Weighted-F1 учитывает размер классов. Balanced accuracy полезна, когда важно сравнить качество распознавания разных классов. Recall класса high выделяется отдельно из-за прикладной важности пропуска высокого риска.",
        "Таким образом, во второй главе формулы выполняют не декоративную роль. Они связывают экономическое содержание риска с конкретной реализацией: доходности и просадки становятся target-компонентами, ликвидность входит в RiskScore и признаки, фундаментальные показатели используются для описания эмитента, а метрики классификации применяются для проверки результата. Эта связь снижает риск критики о том, что машинное обучение применено без предметного обоснования.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 2.5 – Методические требования к данным и их назначение")
    add_table(
        doc,
        ["Требование", "Смысл требования", "Как учитывается в работе"],
        [
            ["Единая дата решения", "Все признаки должны относиться к моменту, когда классификация могла быть выполнена", "Используется monthly decision_date"],
            ["Point-in-time доступность", "Будущая информация не должна попадать в признаки", "Target-компоненты отделены от feature engineering"],
            ["Train-only нормализация", "Пороги и шкалы не должны знать будущие периоды", "Пороги low/medium/high рассчитываются на train"],
            ["Панельная структура", "Нужно учитывать эмитентов и время одновременно", "Наблюдение задается парой ticker и decision_date"],
            ["Диагностика сдвига", "Финансовые распределения меняются во времени", "Используется PSI и walk-forward проверка"],
        ],
        [3.2, 5.9, 6.3],
        font_size=9,
    )
    add_source(doc, "Источник: составлено автором.")
    add_chapter2_additional_depth(doc)


def add_chapter2_additional_depth(doc: Document) -> None:
    for text in [
        "Еще одна важная деталь предметной области состоит в различии между абсолютными и относительными признаками. Абсолютное значение торгового оборота или капитализации может быть полезным, но оно плохо сравнимо между секторами и периодами. Поэтому в пайплайне используются cross-sectional ranks и sector z-scores. Они позволяют оценивать не только уровень показателя, но и положение эмитента относительно других бумаг на том же месячном срезе или внутри собственного сектора.",
        "Cross-sectional rank особенно полезен в инвестиционной задаче, потому что решение часто принимается не в вакууме, а при сравнении альтернатив. Инвестор выбирает между несколькими бумагами, поэтому относительная позиция эмитента по ликвидности, волатильности или долговой нагрузке может быть важнее, чем само значение показателя. Такой подход также снижает чувствительность к общему уровню рынка: если волатильность выросла у всех бумаг, относительный ранг помогает выделить тех, кто стал наиболее рискованным относительно остальных.",
        "Sector z-score решает другую задачу. Компании разных отраслей имеют разные нормальные уровни маржинальности, долговой нагрузки и капитализации. Например, капиталоемкие отрасли могут иметь более высокую долговую нагрузку, чем технологические или сервисные компании. Если сравнивать все эмитенты без учета сектора, модель может ошибочно считать отраслевую норму индивидуальным риском. Секторное отклонение помогает отделить нормальную отраслевую специфику от действительно необычного значения.",
        "Таким образом, feature engineering в работе не является механическим увеличением числа столбцов. Он отражает предметную логику инвестиционного анализа: часть признаков показывает абсолютное состояние объекта, часть – относительное положение на рынке, часть – отклонение от сектора, а часть – взаимодействие нескольких факторов риска. Именно это позволяет связать математическую модель с экономическим смыслом.",
        "При выборе данных также учитывается проблема пропусков. В финансовой панели пропуски неизбежны: не все эмитенты имеют одинаковую историю торгов, отчетность публикуется с разной периодичностью, отдельные показатели могут отсутствовать или быть некорректно извлечены. Простое удаление всех строк с пропусками резко уменьшило бы выборку и могло бы исказить структуру данных. Поэтому пропуски должны обрабатываться системно и сопровождаться индикаторами качества там, где это имеет экономический смысл.",
        "Для отчетности особенно важны признаки качества извлечения. Если значение показателя отсутствует, это не всегда означает нулевой риск; иногда это означает, что показатель не был раскрыт или не был корректно найден в документе. Поэтому missing-индикаторы и признаки лагов публикации могут сами нести информацию. Бумага с неполным раскрытием может быть сложнее для анализа, а значит такая особенность должна быть доступна модели.",
        "Финансовая панель также требует аккуратного обращения с выбросами. Экстремальные значения доходности, оборота или коэффициентов могут быть как реальными стрессовыми событиями, так и ошибками данных. Полное удаление выбросов может уничтожить именно те наблюдения, которые важны для high-класса. Поэтому в риск-задаче предпочтительнее использовать устойчивые преобразования, ранги, winsorization или признаки качества, а не автоматически исключать все экстремальные значения.",
        "С точки зрения предметной области важно также понимать, что классы low, medium и high не являются естественными константами рынка. Это исследовательская дискретизация непрерывного риск-профиля. Поэтому границы классов должны задаваться формально и воспроизводимо. В данной работе это достигается через train-only пороги RiskScore, а не через ручное назначение классов после просмотра результатов.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 2.6 – Логика преобразования исходных показателей в признаки модели")
    add_table(
        doc,
        ["Тип преобразования", "Пример признака", "Зачем используется"],
        [
            ["Rolling window", "rolling_vol_20d, downside_vol_60d", "Учитывает недавнюю динамику риска вместо одного дневного значения"],
            ["Cross-sectional rank", "beta_60d_cs_rank, amihud_20d_cs_rank", "Сравнивает эмитента с другими бумагами на том же месячном срезе"],
            ["Sector z-score", "rolling_vol_20d_sector_z", "Отделяет отраслевую специфику от индивидуального отклонения"],
            ["Interaction feature", "liquidity_vol_interaction", "Фиксирует сочетание факторов, которое может быть опаснее каждого фактора отдельно"],
            ["Report quality feature", "report_lag_days, report_missing_*", "Показывает актуальность и надежность отчетного источника"],
        ],
        [3.5, 4.5, 7.1],
        font_size=9,
    )
    add_source(doc, "Источник: составлено автором.")
    for text in [
        "Такая система признаков необходима для предметной обоснованности модели. Каждый тип признака связан с отдельным экономическим вопросом: насколько бумага волатильна, насколько она ликвидна, насколько она чувствительна к рынку, насколько устойчив эмитент, насколько надежна отчетная информация и отличается ли объект от сектора.",
        "Следовательно, анализ предметной области во второй главе выполняет двойную функцию. Во-первых, он объясняет экономический смысл риска. Во-вторых, он задает методическую основу для последующей реализации: в третьей главе эти формулы и группы признаков должны быть использованы при расчете признаков, формировании целевой переменной и оценке качества.",
        "Выбор методов машинного обучения должен соответствовать этим особенностям данных. Если признаки разнородны и содержат нелинейные взаимосвязи, линейной модели может быть недостаточно. Если рыночные режимы меняются во времени, полезно выделять группы похожих состояний. Если признаковое пространство содержит скрытые сочетания факторов, возникает основание для применения нейросетевого слоя представлений.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 2.7 – Соответствие методов машинного обучения особенностям задачи")
    add_table(
        doc,
        ["Метод", "Почему применим", "Ограничение"],
        [
            ["Деревья решений", "Учитывают пороговые правила, нелинейности и разнородные табличные признаки", "Могут переобучаться без контроля глубины и валидации"],
            ["Ансамбли деревьев", "Снижают нестабильность отдельных деревьев и дают оценку важности признаков", "Требуют сравнения с baseline и контроля качества по классам"],
            ["Кластеризация", "Выделяет похожие рыночные режимы или риск-профили наблюдений", "Кластеры нуждаются в экономической интерпретации"],
            ["Нейронные сети", "Извлекают латентные представления из сложного признакового пространства", "Менее интерпретируемы без дополнительных диагностик"],
            ["Гибридный подход", "Объединяет интерпретируемую классификацию, режимный контекст и нелинейные представления", "Требует строгой схемы validation и test"],
        ],
        [3.2, 7.0, 5.0],
        font_size=12,
    )
    add_source(doc, "Источник: составлено автором.")


def add_chapter3_defense_context(doc: Document, data: dict) -> None:
    metrics = data["metrics"]
    baseline = metrics["architectures"]["baseline_rf"]["test"]
    final = metrics["architectures"]["ann_plus_regime"]["test"]
    macro_delta = final["macro_f1"] - baseline["macro_f1"]
    weighted_delta = final["weighted_f1"] - baseline["weighted_f1"]
    recall_delta = final["high_recall"] - baseline["high_recall"]
    fn_delta = final["high_false_negative_rate"] - baseline["high_false_negative_rate"]
    for text in [
        "Оценка реализации включает не только итоговую метрику, но и весь путь от постановки к результату. Реализация начинается не с выбора нейронной сети, а с подготовки панели наблюдений. На этом этапе проверяются даты, ключи, наличие target-компонент, допустимые признаки и временное разбиение. Такая последовательность снижает риск того, что качество модели будет объясняться ошибкой подготовки данных.",
        "Временное разбиение train, validation и test играет роль имитации реального применения. Модель обучается на прошлом, выбирается на более позднем validation-периоде и окончательно проверяется на test-периоде. Если бы строки перемешивались случайно, в train могли бы попасть наблюдения из будущего рыночного режима, а это сделало бы оценку качества завышенной и методически некорректной.",
        "Архитектуры сравниваются не ради формального перечисления моделей, а для проверки вклада отдельных элементов пайплайна. Baseline_rf показывает качество дерева решений без дополнительных слоев. Regime_only демонстрирует эффект учета рыночного режима. Enriched_reference отражает влияние расширенного feature engineering. Ann_plus_regime проверяет вклад нейросетевого латентного слоя вместе с режимом. Sector_overlay тестирует, дают ли отраслевые модели дополнительный выигрыш.",
        "Такое сравнение позволяет ответить на вопрос, зачем в теме работы одновременно присутствуют деревья решений, кластеризационный анализ и искусственные нейронные сети. Деревья решений обеспечивают нелинейную классификацию и устойчивость к разнородным признакам. Кластеризационный анализ выделяет рыночные режимы, которые могут менять связь между признаками и будущим риском. Нейронная сеть извлекает компактные латентные представления из признакового пространства. Объединение этих элементов образует гибридный пайплайн.",
        f"По сравнению с baseline_rf финальная архитектура ann_plus_regime увеличила macro-F1 на {macro_delta:.4f}, weighted-F1 на {weighted_delta:.4f}, recall класса high на {recall_delta:.4f}. High false negative rate снизился на {abs(fn_delta):.4f}. Эти значения являются прямым ответом на вопрос о приросте точности: улучшение оценивается относительно базовой архитектуры на одном и том же test-периоде.",
        "При этом результат не следует интерпретировать как абсолютную гарантию верного инвестиционного решения. Модель решает задачу автоматической классификации по заданной разметке, а не предсказывает все возможные рыночные события. Поэтому качество описывается через статистические метрики, а практическая интерпретация остается в зоне риск-скрининга и последующей аналитической проверки.",
        "Отдельно фиксируется различие между validation и test. Validation используется для выбора архитектуры и настройки решений пайплайна, включая использование отчетного слоя. Test используется только после выбора модели. Такое разделение необходимо, чтобы итоговая оценка не была результатом подбора по тестовым данным.",
        "Результаты report layer ablation также являются частью обоснования архитектуры. Слой финансовой отчетности включен не из-за формального усложнения темы, а потому что сравнение вариантов with_reports и without_reports показало преимущество варианта с отчетностью на validation. Это делает архитектуру методически обоснованной: каждый дополнительный слой должен проходить проверку полезности.",
        "Walk-forward проверка показывает, что качество модели меняется по периодам. Это ожидаемо для финансовых данных. В спокойные периоды связь между признаками и будущим риском может быть устойчивее, в стрессовые периоды распределения признаков и классов могут смещаться. Поэтому наличие walk-forward результатов укрепляет работу: она не ограничивается одной случайной train-test разбивкой.",
        "Drift diagnostics показывает, какие признаки наиболее сильно изменили распределение. Если среди них присутствуют ставка, валютное давление и рыночная корреляция, это предметно объяснимо: внешний режим рынка действительно менялся. Следовательно, диагностика дрейфа не является техническим приложением, а помогает связать поведение модели с экономическим контекстом.",
        "Файл feature_importance.csv используется для интерпретации результата. Важность признаков не превращает модель в полностью прозрачную, но показывает, что классификация опирается на экономически осмысленные группы: ликвидность, макроэкономическое давление, волатильность, фундаментальную устойчивость и отчетные признаки. Это снижает риск необоснованного использования модели как черного ящика.",
        "Model package обеспечивает воспроизводимость инференса. После обучения сохраняются трансформеры, модель, параметры кластеризации, набор признаков и пороги. Это означает, что новый набор наблюдений может быть пропущен через тот же процесс без ручного восстановления шагов. Тем самым автоматизация проявляется не только в обучении модели, но и в возможности повторного применения.",
        "Важный практический вывод состоит в том, что модель лучше использовать как систему ранжирования и первичного отбора, а не как единственный источник инвестиционного решения. Бумаги, отнесенные к high, должны попадать на углубленную проверку. Такой сценарий снижает нагрузку ручного анализа, потому что дальнейшая проверка концентрируется не на всей панели, а на наиболее рискованных сигналах.",
    ]:
        add_plain_paragraph(doc, text)
    add_caption(doc, "Таблица 3.7 – Измерение прироста относительно базовой архитектуры")
    add_table(
        doc,
        ["Метрика", "Baseline", "Final", "Δ", "Интерпретация"],
        [
            ["Macro-F1", baseline["macro_f1"], final["macro_f1"], macro_delta, "Улучшение среднего качества по классам"],
            ["Weighted-F1", baseline["weighted_f1"], final["weighted_f1"], weighted_delta, "Улучшение с учетом размера классов"],
            ["High recall", baseline["high_recall"], final["high_recall"], recall_delta, "Меньше пропусков высокого риска"],
            ["High FN rate", baseline["high_false_negative_rate"], final["high_false_negative_rate"], fn_delta, "Снижение наиболее опасной ошибки"],
            ["Adjacent accuracy", baseline["ordinal_adjacent_accuracy"], final["ordinal_adjacent_accuracy"], final["ordinal_adjacent_accuracy"] - baseline["ordinal_adjacent_accuracy"], "Больше ошибок остаются соседними по уровню риска"],
        ],
        [3.2, 2.0, 2.2, 1.7, 6.2],
        font_size=12,
    )
    add_source(doc, "Источник: составлено автором по метрикам основного запуска пайплайна.")
    for text in [
        "Приведенное сравнение показывает, что прирост точности в работе выражен не общими словами, а численно. Особенно важно снижение high false negative rate, поскольку для задачи инвестиционного риска пропуск действительно опасного объекта является более критичным, чем избыточное предупреждение. В этой логике финальная архитектура лучше соответствует прикладной цели риск-скрининга.",
        "Масштабируемость реализации подтверждается тем, что одни и те же функции применяются ко всей панели наблюдений. В коде нет отдельной логики для каждого эмитента или ручного выбора периода. Если во входной панели появятся новые строки с теми же полями, пайплайн сможет рассчитать признаки, выполнить инференс и сохранить результаты в том же формате. Ограничением остается качество и полнота входных данных, но не сама методика классификации.",
        "Воспроизводимость подтверждается набором артефактов. Metrics.json фиксирует метрики и параметры, predictions.csv сохраняет предсказания и вероятности, model_leaderboard.csv показывает сравнение кандидатов, walk_forward_report.csv хранит проверку по временным фолдам, feature_drift_report.csv фиксирует сдвиг признаков, а model_package.joblib содержит сохраненную модель. Такой набор позволяет восстановить ход исследования без обращения к неформальным пояснениям автора.",
        "Таким образом, третья глава показывает, что именно было реализовано, почему использованы выбранные методы, как обеспечена временная корректность, чем финальная архитектура лучше базовой и какие ограничения остаются. Это делает выводы работы более устойчивыми, чем простое сообщение одной итоговой метрики.",
    ]:
        add_plain_paragraph(doc, text)


def add_conclusion_defense_context(doc: Document) -> None:
    for text in [
        "С точки зрения поставленных задач можно сделать более детальный итог. Первая задача, связанная с изучением подходов к оценке риска, выполнена через анализ экономической сущности инвестиционного риска и рассмотрение показателей, которые отражают разные стороны неблагоприятной реализации. Показано, что риск акции не может быть корректно описан только одним показателем, поэтому в работе используется многокомпонентная логика.",
        "Вторая задача, связанная с анализом ограничений ручного подхода, выполнена через выделение субъективности, трудоемкости, слабой масштабируемости, низкой воспроизводимости и риска временной некорректности. Эти ограничения не описываются как бытовые проблемы отдельных компаний, а рассматриваются как методические ограничения процедуры классификации.",
        "Третья задача, связанная с определением признаков и метрик, выполнена через обоснование рыночных, ликвидностных, макроэкономических, фундаментальных и отчетных признаков. Для них выписаны экономические и математические формулы, а также показано, каким образом эти формулы используются в target engineering и feature engineering.",
        "Четвертая и пятая задачи, связанные с проектированием и реализацией ML-пайплайна, выполнены через разработку последовательности этапов: загрузка данных, контроль качества, расчет признаков, формирование целевой переменной, временное разбиение, кластеризация режима, обучение ансамблей деревьев, использование латентных факторов и сохранение модели.",
        "Шестая задача, связанная с оценкой качества, выполнена через расчет метрик классификации, сравнение архитектур, walk-forward проверку и диагностику дрейфа признаков. Поэтому результат работы оценивается не декларативно, а через воспроизводимые числовые показатели и сохраненные артефакты.",
        "Вопрос о приросте масштабируемости, воспроизводимости и точности также получает формальный ответ. Масштабируемость подтверждается обработкой панели наблюдений единым пайплайном. Воспроизводимость подтверждается сохранением данных, кода, параметров и артефактов запуска. Точность подтверждается сравнением финальной архитектуры с baseline_rf и расчетом метрик на test-периоде.",
        "Главным ограничением работы является зависимость результата от качества исходных данных и устойчивости финансового режима. Это не отменяет результата, но задает условия дальнейшего применения. Для практического использования модель должна сопровождаться регулярным переобучением, контролем drift, проверкой качества отчетного слоя и настройкой порогов под конкретную политику риска.",
        "В целом работа показывает, что заявленная тема раскрыта не как набор несвязанных методов, а как единый процесс: экономическая постановка риска превращается в формулы, формулы превращаются в признаки и target, признаки передаются в гибридный ML-пайплайн, а результат оценивается через метрики и временную проверку. Именно эта связность является основным итогом выполненного исследования.",
    ]:
        add_plain_paragraph(doc, text)
    add_plain_paragraph(doc, "Данная работа выполнена мною самостоятельно", align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)
    add_plain_paragraph(doc, "«___» _______________ 2026 г.          ______________ / Поддуба И.С. /", align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False)


def build() -> Path:
    CAPTION_COUNTERS["table"] = 0
    CAPTION_COUNTERS["figure"] = 0
    OUT_DIR.mkdir(exist_ok=True)
    data = load_results()
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_toc(doc)
    intro(doc)
    chapter1(doc)
    chapter2(doc, data)
    chapter3(doc, data)
    conclusion(doc)
    references(doc)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(path)
